"""docx_writer 单元测试（C2 / F15）

覆盖 C1 引入的能力：
- ![cap](path) 行渲染为 inline image / 占位符
- setup_multilevel_list 创建 abstractNum + Heading 1-5 numPr 绑定
- add_toc_field 插入 TOC 域
- validate_heading_hierarchy 跳级检测
- strip_numbering_prefix 兜底
- setup_styles 段距 / 首行缩进 / 行距生效
"""

import os
import re
import sys
import tempfile
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT / 'skills' / 'taw' / 'tools'))

import docx_writer
from docx.oxml.ns import qn


def _dummy_png(path: str):
    """生成 1x1 PNG。"""
    try:
        from PIL import Image
        Image.new('RGB', (1, 1), color=(0, 0, 0)).save(path, format='PNG')
    except ImportError:
        import struct, zlib

        def chunk(t, d):
            return struct.pack('>I', len(d)) + t + d + struct.pack('>I', zlib.crc32(t + d))

        ihdr = struct.pack('>IIBBBBB', 1, 1, 8, 0, 0, 0, 0)
        idat = zlib.compress(b'\x00\x00')
        with open(path, 'wb') as f:
            f.write(b'\x89PNG\r\n\x1a\n')
            f.write(chunk(b'IHDR', ihdr))
            f.write(chunk(b'IDAT', idat))
            f.write(chunk(b'IEND', b''))


# ── Test 1: write_markdown 解析 ![]() ─────────────

def test_write_markdown_parses_image_ref(tmp_path):
    img = tmp_path / 'pic.png'
    _dummy_png(str(img))
    doc = docx_writer.create_document()
    docx_writer.write_markdown(doc, f'前导段。\n\n![图 1.3-1：示意图]({img})\n\n后续段。')
    # 找出 inline 图片元素
    img_count = sum(1 for r in doc.element.iter(qn('w:drawing')))
    assert img_count >= 1, f'图片未被嵌入（drawing count={img_count}）'
    # caption 段：找含 '图 1.3-1' 的段落
    caps = [p for p in doc.paragraphs if '图 1.3-1' in p.text]
    assert caps, 'caption 段未生成'


def test_write_markdown_image_missing_falls_back_to_placeholder():
    doc = docx_writer.create_document()
    docx_writer.write_markdown(doc, '![图 1.3-2：缺图](/nonexistent/path.png)')
    # 找占位符段
    placeholders = [p for p in doc.paragraphs if '图片占位符' in p.text]
    assert placeholders, '占位符段未生成'


# ── Test 2: setup_multilevel_list ─────────────────

def test_setup_multilevel_list_creates_numbering_xml():
    doc = docx_writer.create_document()  # 内部已调 setup_multilevel_list
    numbering_part = doc.part.numbering_part
    abs_nums = numbering_part.element.findall(qn('w:abstractNum'))
    assert abs_nums, 'abstractNum 未创建'
    # 找 5 个 lvl
    abs_num = abs_nums[-1]
    lvls = abs_num.findall(qn('w:lvl'))
    assert len(lvls) == 5, f'期望 5 个 lvl，实际 {len(lvls)}'


def test_setup_multilevel_list_binds_headings_1_to_5():
    doc = docx_writer.create_document()
    for h in range(1, 6):
        style = doc.styles[f'Heading {h}']
        pPr = style.element.find(qn('w:pPr'))
        assert pPr is not None, f'Heading {h} 缺 pPr'
        numPr = pPr.find(qn('w:numPr'))
        assert numPr is not None, f'Heading {h} 未绑定多级列表'


# ── Test 3: add_toc_field ─────────────────────────

def test_add_toc_field_inserts_field_code():
    doc = docx_writer.create_document()
    docx_writer.add_toc_field(doc, levels=4)
    # 找含 'TOC' 域指令的元素
    instr_texts = [el.text for el in doc.element.iter(qn('w:instrText'))]
    toc_instrs = [t for t in instr_texts if t and 'TOC' in t]
    assert toc_instrs, 'TOC 域未插入'
    assert any('1-4' in t for t in toc_instrs), 'TOC levels 参数错'


# ── Test 4: validate_heading_hierarchy ────────────

def test_validate_heading_hierarchy_passes_aligned_doc():
    doc = docx_writer.create_document()
    docx_writer.write_markdown(doc, '# 一级\n## 二级\n### 三级')
    ok, errs = docx_writer.validate_heading_hierarchy(doc, [
        {'numbering': '1', 'title': '一级', 'depth': 1},
        {'numbering': '1.1', 'title': '二级', 'depth': 2},
        {'numbering': '1.1.1', 'title': '三级', 'depth': 3},
    ])
    assert ok, f'对齐 doc 应通过，errs={errs}'


def test_validate_heading_hierarchy_detects_level_mismatch():
    doc = docx_writer.create_document()
    # 大纲期望 H3，doc 写成 H1（subagent bug 模拟）
    docx_writer.write_markdown(doc, '# 招标方需求理解')
    ok, errs = docx_writer.validate_heading_hierarchy(doc, [
        {'numbering': '1.2.1', 'title': '招标方需求理解', 'depth': 3},
    ])
    assert not ok
    assert any('Heading 3' in e and 'Heading 1' in e for e in errs), errs


# ── Test 5: strip_numbering_prefix ────────────────

def test_strip_numbering_prefix():
    assert docx_writer.strip_numbering_prefix('1.3.1 招标方需求理解') == '招标方需求理解'
    assert docx_writer.strip_numbering_prefix('1. 技术部分') == '技术部分'
    assert docx_writer.strip_numbering_prefix('  1.2  二级标题  ') == '二级标题'
    assert docx_writer.strip_numbering_prefix('无编号标题') == '无编号标题'


def test_add_heading_cn_strips_residual_numbering():
    """若 subagent 残留编号写进 heading text，运行时应剥除"""
    doc = docx_writer.create_document()
    docx_writer.add_heading_cn(doc, '1.3.1 招标方需求理解', level=3)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    assert '招标方需求理解' in headings, headings
    assert '1.3.1 招标方需求理解' not in headings, '编号未剥除'


# ── Test 6: setup_styles 行距/段距/缩进 ───────────

def test_setup_styles_normal_paragraph_format():
    from docx.shared import Pt
    doc = docx_writer.create_document()
    pf = doc.styles['Normal'].paragraph_format
    assert pf.line_spacing == 1.5, f'行距 {pf.line_spacing}'
    assert pf.first_line_indent == Pt(24), f'首行缩进 {pf.first_line_indent}'
    assert pf.space_before == Pt(6), f'段前 {pf.space_before}'
    assert pf.space_after == Pt(6), f'段后 {pf.space_after}'


def test_setup_styles_heading_paragraph_format():
    from docx.shared import Pt
    doc = docx_writer.create_document()
    h2_pf = doc.styles['Heading 2'].paragraph_format
    assert h2_pf.line_spacing == 2.0, f'H2 行距 {h2_pf.line_spacing}'
    assert h2_pf.first_line_indent == Pt(0), f'H2 不应缩进（{h2_pf.first_line_indent}）'


# ── Test 7: write_markdown level 0 clamp ─────────

def test_write_markdown_does_not_render_level_zero():
    """`#` 走 add_heading(text, 1)，绝不调 add_heading(text, 0)（Title 样式）"""
    doc = docx_writer.create_document()
    docx_writer.write_markdown(doc, '# 顶层')
    titles = [p.text for p in doc.paragraphs if p.style.name == 'Title']
    headings = [p.text for p in doc.paragraphs if p.style.name == 'Heading 1']
    assert '顶层' in headings, headings
    assert '顶层' not in titles, '不应渲染为 Title'
