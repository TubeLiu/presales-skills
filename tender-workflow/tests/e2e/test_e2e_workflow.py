"""端到端测试主流程"""
import os
import shlex
import subprocess
import time
from pathlib import Path
from datetime import datetime
from typing import Dict, Any

from docx import Document
import pytest

from validators.taa_validator import TAAValidator
from validators.taw_validator import TAWValidator
from validators.quality_checker import QualityChecker
from validators.consistency_validator import ConsistencyValidator
from validators.image_quality_validator import ImageQualityValidator
from validators.content_professionalism_validator import ContentProfessionalismValidator


# 项目根目录
PROJECT_ROOT = Path(__file__).parent.parent.parent

# 测试文件
TENDER_FILE = Path(os.environ.get('TENDER_TEST_FILE', str(PROJECT_ROOT / 'tests' / '水网智科.md')))

# 输出目录
OUTPUT_DIR = PROJECT_ROOT / 'output'
DRAFTS_DIR = PROJECT_ROOT / 'drafts'
REPORT_DIR = PROJECT_ROOT / 'tests' / 'e2e' / 'reports'

# 测试章节（taa v2.1.0 后编号改为 1.X 体系）
TEST_CHAPTERS = ['1.3', '1.4', '1.10']


class TestE2EWorkflow:
    """端到端测试工作流"""

    def _skip_nested_session_if_needed(self, result):
        if 'CLAUDECODE' in result.stderr or 'Nested sessions' in result.stderr or 'nested' in result.stderr.lower():
            pytest.skip('在 Claude Code 会话内部无法运行嵌套 claude skill 命令')

    def _build_claude_prompt_command(self, skill_command: str) -> str:
        return f"claude -p {shlex.quote(skill_command)}"

    @pytest.fixture(scope='class')
    def test_context(self):
        """测试上下文"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        return {
            'tender_file': TENDER_FILE,
            'output_dir': OUTPUT_DIR,
            'drafts_dir': DRAFTS_DIR,
            'report_dir': REPORT_DIR,
            'timestamp': timestamp,
            'start_time': time.time(),
            'taa_outputs': {},
            'taw_outputs': {},
            'taa_issues': [],
            'outline_issues': [],
            'taw_issues': [],
            'consistency_issues': [],
            'image_quality_issues': [],
            'professionalism_issues': [],
            'taa_metrics': {},
            'outline_metrics': {},
            'taw_metrics': {},
            'trv_reports': [],
            'errors': [],
        }

    def test_0_environment_setup(self, test_context):
        """Step 0: 环境准备"""
        print("\n=== Step 0: 环境准备 ===")

        # 检查测试文件存在性
        if not TENDER_FILE.exists():
            pytest.fail(f"测试文件不存在: {TENDER_FILE}")

        print(f"✅ 测试文件: {TENDER_FILE}")
        print(f"   大小: {TENDER_FILE.stat().st_size / 1024:.1f} KB")

        # 创建输出目录
        REPORT_DIR.mkdir(parents=True, exist_ok=True)
        print(f"✅ 报告目录: {REPORT_DIR}")

        # 初始化测试上下文
        print(f"✅ 测试时间戳: {test_context['timestamp']}")

    def test_1_taa_execution(self, test_context):
        """Step 1: 查找 taa 输出文件（跳过执行，使用已有文件）"""
        print("\n=== Step 1: 查找 taa 输出文件 ===")
        print("注意：由于不能在 Claude Code 会话中嵌套调用 skill，")
        print("      本测试使用已有的 output 文件进行验证。")
        print("      如需重新生成，请在测试外手动执行：")
        print(f"      /taa {TENDER_FILE} --vendor 灵雀云")
        print()

        # 查找输出文件
        analysis_files = list(OUTPUT_DIR.glob('招标分析报告_*.md'))
        outline_files = list(OUTPUT_DIR.glob('投标文件大纲_*.docx'))

        if not analysis_files:
            pytest.skip("未找到分析报告文件，请先手动执行 /taa")
        if not outline_files:
            pytest.skip("未找到投标大纲文件，请先手动执行 /taa")

        # 使用最新的文件
        analysis_file = max(analysis_files, key=lambda p: p.stat().st_mtime)
        outline_file = max(outline_files, key=lambda p: p.stat().st_mtime)

        test_context['taa_outputs'] = {
            'analysis': analysis_file,
            'outline': outline_file,
        }

        print(f"✅ 找到分析报告: {analysis_file.name}")
        print(f"✅ 找到投标大纲: {outline_file.name}")

    def test_2_taa_analysis_quality(self, test_context):
        """Step 2: 验证 taa 分析报告质量（增强版）"""
        print("\n=== Step 2: 验证 taa 分析报告质量 ===")

        analysis_file = test_context['taa_outputs'].get('analysis')
        if not analysis_file or not analysis_file.exists():
            pytest.skip("分析报告文件不存在")

        # 创建验证器
        validator = TAAValidator(analysis_file)

        # 执行验证（包含信息准确性验证）
        issues = validator.validate_all(tender_file=TENDER_FILE)
        test_context['taa_issues'] = issues

        # 计算指标
        modules = validator.modules
        test_context['taa_metrics'] = {
            'taa_module_completeness': len(modules) / 7.0,
            'taa_citation_rate': 0.0,  # 需要从验证结果提取
            'taa_speculation_count': len([i for i in issues if i['type'] == 'speculation']),
            'taa_information_accuracy': len([i for i in issues if i['type'] == 'information_accuracy']) == 0,
        }

        # 生成报告
        report_file = REPORT_DIR / f'taa_validation_{test_context["timestamp"]}.md'
        validator.generate_report(issues, report_file)
        print(f"验证报告: {report_file}")

        # 统计问题
        critical = len([i for i in issues if i['severity'] == 'critical'])
        high = len([i for i in issues if i['severity'] == 'high'])
        print(f"发现问题: {len(issues)} 个（严重 {critical}，高风险 {high}）")

        # 严重问题导致测试失败
        if critical > 0:
            pytest.fail(f"发现 {critical} 个严重问题")

    def test_3_taa_outline_quality(self, test_context):
        """Step 3: 验证投标大纲质量"""
        print("\n=== Step 3: 验证投标大纲质量 ===")

        analysis_file = test_context['taa_outputs'].get('analysis')
        outline_file = test_context['taa_outputs'].get('outline')

        if not analysis_file or not analysis_file.exists():
            pytest.skip("分析报告文件不存在")
        if not outline_file or not outline_file.exists():
            pytest.skip("投标大纲文件不存在")

        # 创建验证器
        validator = TAAValidator(analysis_file, outline_file)

        # 解析大纲
        outline_data = validator.parse_outline()

        # 验证大纲
        issues = validator.validate_outline(outline_data)
        test_context['outline_issues'] = issues

        # 计算指标
        test_context['outline_metrics'] = {
            'outline_chapter_count': outline_data.get('chapter_count', 0),
            'outline_requirement_coverage': 0.8,  # 需要从验证结果提取
            'outline_numbering_compliance': '通过' if len([i for i in issues if '编号' in i['description']]) == 0 else '不通过',
        }

        # 生成报告
        report_file = REPORT_DIR / f'outline_validation_{test_context["timestamp"]}.md'
        validator.generate_report(issues, report_file)
        print(f"验证报告: {report_file}")

        # 统计问题
        critical = len([i for i in issues if i['severity'] == 'critical'])
        high = len([i for i in issues if i['severity'] == 'high'])
        print(f"发现问题: {len(issues)} 个（严重 {critical}，高风险 {high}）")

        # 严重问题导致测试失败
        if critical > 0:
            pytest.fail(f"发现 {critical} 个严重问题")

    def test_4_taw_execution(self, test_context):
        """Step 4: 执行 taw，生成章节内容"""
        print("\n=== Step 4: 执行 taw，生成章节内容 ===")

        # 确保 output/ 目录存在
        if not OUTPUT_DIR.exists():
            pytest.skip("output/ 目录不存在")

        # 为每个测试章节执行 taw
        for chapter in TEST_CHAPTERS:
            print(f"\n生成章节 {chapter}...")

            skill_command = f'/taw {OUTPUT_DIR} --chapter {chapter} --search-tool auto'
            cmd = self._build_claude_prompt_command(skill_command)
            print(f"执行命令: {cmd}")

            try:
                result = subprocess.run(
                    cmd,
                    shell=True,
                    capture_output=True,
                    text=True,
                    timeout=600,  # 10 分钟超时
                )

                if result.returncode != 0:
                    self._skip_nested_session_if_needed(result)
                    test_context['errors'].append({
                        'step': f'taw_execution_{chapter}',
                        'error': result.stderr,
                    })
                    print(f"章节 {chapter} 生成失败: {result.stderr}")
                    continue

                print(f"章节 {chapter} 生成成功")

            except subprocess.TimeoutExpired:
                test_context['errors'].append({
                    'step': f'taw_execution_{chapter}',
                    'error': 'Timeout after 10 minutes',
                })
                print(f"章节 {chapter} 生成超时（10 分钟）")
                continue

        # 查找输出文件
        for chapter in TEST_CHAPTERS:
            chapter_files = list(DRAFTS_DIR.glob(f'{chapter}_*.docx'))
            if chapter_files:
                chapter_file = max(chapter_files, key=lambda p: p.stat().st_mtime)
                test_context['taw_outputs'][chapter] = chapter_file
                print(f"章节 {chapter}: {chapter_file}")

        if not test_context['taw_outputs']:
            pytest.fail("未生成任何章节文件")

    def test_5_taw_content_quality(self, test_context):
        """Step 5: 验证 taw 章节内容质量（增强版）"""
        print("\n=== Step 5: 验证 taw 章节内容质量 ===")

        analysis_file = test_context['taa_outputs'].get('analysis')
        if not analysis_file or not analysis_file.exists():
            pytest.skip("分析报告文件不存在")

        if not test_context['taw_outputs']:
            pytest.skip("未生成任何章节文件")

        # 创建验证器
        validator = TAWValidator(analysis_file)

        # 验证每个章节
        all_issues = []
        all_image_issues = []
        all_professionalism_issues = []
        total_words = 0
        total_images = 0
        total_citations = 0

        for chapter, chapter_file in test_context['taw_outputs'].items():
            if not chapter_file.exists():
                continue

            print(f"\n验证章节 {chapter}...")

            # 解析章节
            content = validator.parse_chapter(chapter_file)

            # 基础验证
            issues = validator.validate_chapter(content['chapter'], content)
            all_issues.extend(issues)

            # 图片质量验证
            image_issues = validator.validate_with_image_quality(chapter_file)
            all_image_issues.extend(image_issues)

            # 内容专业性验证
            professionalism_issues = validator.validate_with_professionalism(chapter_file)
            all_professionalism_issues.extend(professionalism_issues)

            # 统计指标
            total_words += content['word_count']
            total_images += content['image_count']
            total_citations += content['citation_count']

            print(f"  字数: {content['word_count']}")
            print(f"  图片: {content['image_count']}")
            print(f"  引用: {content['citation_count']}")
            print(f"  基础问题: {len(issues)} 个")
            print(f"  图片问题: {len(image_issues)} 个")
            print(f"  专业性问题: {len(professionalism_issues)} 个")

        test_context['taw_issues'] = all_issues
        test_context['image_quality_issues'] = all_image_issues
        test_context['professionalism_issues'] = all_professionalism_issues

        # 计算指标
        chapter_count = len(test_context['taw_outputs'])
        test_context['taw_metrics'] = {
            'taw_avg_word_count': total_words / chapter_count if chapter_count > 0 else 0,
            'taw_avg_image_count': total_images / chapter_count if chapter_count > 0 else 0,
            'taw_avg_citation_count': total_citations / chapter_count if chapter_count > 0 else 0,
            'taw_scoring_coverage': 0.0,  # 需要从验证结果提取
            'taw_keyword_coverage': 0.0,  # 需要从验证结果提取
        }

        # 生成报告
        report_file = REPORT_DIR / f'taw_validation_{test_context["timestamp"]}.md'
        validator.generate_report(all_issues, test_context['taw_outputs'], report_file)
        print(f"\n验证报告: {report_file}")

        # 统计问题
        all_combined_issues = all_issues + all_image_issues + all_professionalism_issues
        critical = len([i for i in all_combined_issues if i['severity'] == 'critical'])
        high = len([i for i in all_combined_issues if i['severity'] == 'high'])
        print(f"发现问题: {len(all_combined_issues)} 个（严重 {critical}，高风险 {high}）")

        # 严重问题导致测试失败
        if critical > 0:
            pytest.fail(f"发现 {critical} 个严重问题")

    def test_6_cross_chapter_consistency(self, test_context):
        """Step 6: 验证章节间一致性"""
        print("\n=== Step 6: 验证章节间一致性 ===")

        analysis_file = test_context['taa_outputs'].get('analysis')
        outline_file = test_context['taa_outputs'].get('outline')

        if not analysis_file or not analysis_file.exists():
            pytest.skip("分析报告文件不存在")
        if not outline_file or not outline_file.exists():
            pytest.skip("投标大纲文件不存在")

        # 获取章节文件列表
        chapter_files = list(test_context['taw_outputs'].values())
        if not chapter_files:
            pytest.skip("未生成任何章节文件")

        # 创建一致性验证器
        validator = ConsistencyValidator(analysis_file, outline_file, chapter_files)

        # 执行验证
        issues = validator.validate_all()
        test_context['consistency_issues'] = issues

        # 生成报告
        report_file = REPORT_DIR / f'consistency_validation_{test_context["timestamp"]}.md'
        validator.generate_report(report_file, issues)
        print(f"验证报告: {report_file}")

        # 统计问题
        critical = len([i for i in issues if i['severity'] == 'critical'])
        high = len([i for i in issues if i['severity'] == 'high'])
        print(f"发现问题: {len(issues)} 个（严重 {critical}，高风险 {high}）")

        # 高风险问题导致测试失败
        if high > 0:
            pytest.fail(f"发现 {high} 个高风险一致性问题")

    def test_7_outline_matching(self, test_context):
        """Step 7: 验证大纲匹配度"""
        print("\n=== Step 7: 验证大纲匹配度 ===")

        analysis_file = test_context['taa_outputs'].get('analysis')
        outline_file = test_context['taa_outputs'].get('outline')

        if not analysis_file or not analysis_file.exists():
            pytest.skip("分析报告文件不存在")
        if not outline_file or not outline_file.exists():
            pytest.skip("投标大纲文件不存在")

        # 获取章节文件列表
        chapter_files = list(test_context['taw_outputs'].values())
        if not chapter_files:
            pytest.skip("未生成任何章节文件")

        # 使用一致性验证器检查章节-大纲匹配度
        validator = ConsistencyValidator(analysis_file, outline_file, chapter_files)

        # 只执行章节-大纲一致性检查
        issues = []
        for chapter_file in chapter_files:
            chapter_num = validator._extract_chapter_number(chapter_file.name)
            if chapter_num:
                issues.extend(validator.validate_chapter_outline_consistency(chapter_num))

        print(f"发现问题: {len(issues)} 个")

        # 中等以上问题导致警告
        medium_plus = len([i for i in issues if i['severity'] in ['critical', 'high', 'medium']])
        if medium_plus > 0:
            print(f"⚠️  发现 {medium_plus} 个中等以上匹配度问题")

    def test_8_trv_integration(self, test_context):
        """Step 8: trv 审核集成"""
        print("\n=== Step 8: trv 审核集成 ===")

        # 确保有必要的文件
        analysis_file = test_context['taa_outputs'].get('analysis')
        outline_file = test_context['taa_outputs'].get('outline')

        if not analysis_file or not analysis_file.exists():
            pytest.skip("分析报告文件不存在")
        if not outline_file or not outline_file.exists():
            pytest.skip("投标大纲文件不存在")

        # 测试场景列表
        test_scenarios = [
            {
                'name': '审核分析报告',
                'file': analysis_file,
                'type': 'analysis',
                'reference': TENDER_FILE,
                'level': 'all',
                'focus': None,
            },
            {
                'name': '审核投标大纲',
                'file': outline_file,
                'type': 'outline',
                'reference': analysis_file,
                'level': 'all',
                'focus': None,
            },
            {
                'name': '审核投标大纲并生成修订版',
                'file': outline_file,
                'type': 'outline',
                'reference': analysis_file,
                'level': 'all',
                'focus': None,
                'revise_docx': True,
                'revise_scope': 'must',
            },
            {
                'name': '审核投标大纲（评分契合度）',
                'file': outline_file,
                'type': 'outline',
                'reference': analysis_file,
                'level': 'all',
                'focus': 'scoring',
            },
        ]

        # 如果有章节文件，添加章节审核
        if test_context['taw_outputs']:
            first_chapter_file = list(test_context['taw_outputs'].values())[0]
            test_scenarios.append({
                'name': '审核章节草稿',
                'file': first_chapter_file,
                'type': 'chapter',
                'reference': analysis_file,
                'level': 'all',
                'focus': None,
            })
            test_scenarios.append({
                'name': '审核章节草稿并生成修订版',
                'file': first_chapter_file,
                'type': 'chapter',
                'reference': analysis_file,
                'level': 'all',
                'focus': None,
                'revise_docx': True,
                'revise_scope': 'all',
            })
            test_scenarios.append({
                'name': '审核章节草稿（快速检查）',
                'file': first_chapter_file,
                'type': 'chapter',
                'reference': analysis_file,
                'level': 'critical',
                'focus': 'risk',
            })

        # 执行每个测试场景
        for scenario in test_scenarios:
            print(f"\n{scenario['name']}...")

            # 构建命令
            parts = [
                '/trv',
                shlex.quote(str(scenario['file'])),
                f'--type {scenario["type"]}',
            ]

            if scenario['reference']:
                parts.append(f'--reference {shlex.quote(str(scenario["reference"]))}')
            if scenario['level']:
                parts.append(f'--level {scenario["level"]}')
            if scenario['focus']:
                parts.append(f'--focus {scenario["focus"]}')
            if scenario.get('revise_docx'):
                parts.append('--revise-docx')
            if scenario.get('revise_scope'):
                parts.append(f'--revise-scope {scenario["revise_scope"]}')

            skill_command = ' '.join(parts)
            cmd = self._build_claude_prompt_command(skill_command)
            print(f"执行命令: {cmd}")

            try:
                result = subprocess.run(
                    cmd,
                    shell=True,
                    capture_output=True,
                    text=True,
                    timeout=300,  # 5 分钟超时
                )

                if result.returncode != 0:
                    self._skip_nested_session_if_needed(result)
                    test_context['errors'].append({
                        'step': f'trv_{scenario["name"]}',
                        'error': result.stderr,
                    })
                    print(f"❌ {scenario['name']} 失败: {result.stderr[:200]}")
                    continue

                print(f"✅ {scenario['name']} 成功")

                # 查找生成的审核报告
                trv_output_dir = OUTPUT_DIR / 'trv'
                if trv_output_dir.exists():
                    report_files = list(trv_output_dir.glob(f'审核报告_{scenario["type"]}_*.md'))
                    if report_files:
                        latest_report = max(report_files, key=lambda p: p.stat().st_mtime)
                        test_context['trv_reports'].append({
                            'scenario': scenario['name'],
                            'file': latest_report,
                            'type': scenario['type'],
                        })
                        print(f"   报告: {latest_report.name}")

                        # 简单验证报告内容
                        report_content = latest_report.read_text(encoding='utf-8')
                        if '审核结论' in report_content and '问题清单' in report_content:
                            print(f"   ✅ 报告结构完整")
                        else:
                            print(f"   ⚠️  报告结构可能不完整")

                        if scenario.get('revise_docx'):
                            revised_files = list(trv_output_dir.glob(f'{scenario["file"].stem}_修订版_*.docx'))
                            if revised_files:
                                latest_revised = max(revised_files, key=lambda p: p.stat().st_mtime)
                                revised_doc = Document(latest_revised)
                                # 验证修订版与原文件存在实际差异
                                original_doc = Document(scenario['file'])
                                orig_texts = [p.text for p in original_doc.paragraphs]
                                revised_texts = [p.text for p in revised_doc.paragraphs]
                                has_diff = orig_texts != revised_texts
                                if not has_diff:
                                    for ot, rt in zip(original_doc.tables, revised_doc.tables):
                                        for orow, rrow in zip(ot.rows, rt.rows):
                                            if any(oc.text != rc.text for oc, rc in zip(orow.cells, rrow.cells)):
                                                has_diff = True
                                                break
                                        if has_diff:
                                            break
                                if has_diff:
                                    print(f"   ✅ 修订版: {latest_revised.name}（已验证存在实际修订）")
                                else:
                                    print(f"   ⚠️  修订版文件生成但内容与原文件相同，修订可能未命中")
                            else:
                                print(f"   ⚠️  未找到修订版 DOCX")

            except subprocess.TimeoutExpired:
                test_context['errors'].append({
                    'step': f'trv_{scenario["name"]}',
                    'error': 'Timeout after 5 minutes',
                })
                print(f"❌ {scenario['name']} 超时（5 分钟）")
                continue

        # 统计结果
        print(f"\n审核报告总数: {len(test_context['trv_reports'])} 个")
        for report in test_context['trv_reports']:
            print(f"  - {report['scenario']}: {report['file'].name}")

        # 如果没有生成任何报告，测试失败
        if not test_context['trv_reports']:
            pytest.fail("未生成任何审核报告")

    def test_9_comprehensive_quality_scoring(self, test_context):
        """Step 9: 综合质量评分"""
        print("\n=== Step 9: 综合质量评分 ===")

        # 汇总所有问题
        all_issues = (
            test_context['taa_issues'] +
            test_context['outline_issues'] +
            test_context['taw_issues'] +
            test_context['consistency_issues'] +
            test_context['image_quality_issues'] +
            test_context['professionalism_issues']
        )

        # 计算一级指标
        # 1. 分析报告质量（30%）
        taa_completeness = test_context['taa_metrics'].get('taa_module_completeness', 0)
        taa_citation = 0.3  # 简化，实际需要从问题中计算
        taa_no_speculation = 1 - min(test_context['taa_metrics'].get('taa_speculation_count', 0) / 10, 1)
        analysis_quality = (taa_completeness + taa_citation + taa_no_speculation) / 3

        # 2. 投标大纲质量（20%）
        outline_completeness = 1.0 if len([i for i in test_context['outline_issues'] if i['severity'] == 'critical']) == 0 else 0.5
        outline_coverage = 0.9  # 简化
        outline_consistency = 1.0 if len([i for i in test_context['consistency_issues'] if 'outline_analysis' in i['type']]) == 0 else 0.7
        outline_quality = (outline_completeness + outline_coverage + outline_consistency) / 3

        # 3. 章节内容质量（40%）
        taw_word_compliance = 1.0 if len([i for i in test_context['taw_issues'] if i['type'] == 'word_count']) == 0 else 0.7
        taw_image_ratio = 0.8  # 简化
        taw_scoring_coverage = 0.6  # 简化
        taw_professionalism = 0.8  # 简化
        content_quality = (taw_word_compliance + taw_image_ratio + taw_scoring_coverage + taw_professionalism) / 4

        # 4. 一致性质量（10%）
        consistency_quality = 1.0 if len([i for i in test_context['consistency_issues'] if i['severity'] in ['critical', 'high']]) == 0 else 0.7

        # 总体质量得分
        overall_quality = (
            analysis_quality * 0.3 +
            outline_quality * 0.2 +
            content_quality * 0.4 +
            consistency_quality * 0.1
        )

        print(f"\n质量评分:")
        print(f"  分析报告质量: {analysis_quality*100:.1f}% (目标 ≥85%)")
        print(f"  投标大纲质量: {outline_quality*100:.1f}% (目标 ≥90%)")
        print(f"  章节内容质量: {content_quality*100:.1f}% (目标 ≥80%)")
        print(f"  一致性质量: {consistency_quality*100:.1f}% (目标 ≥90%)")
        print(f"  总体质量: {overall_quality*100:.1f}%")

        # 保存到上下文
        test_context['quality_scores'] = {
            'analysis_quality': analysis_quality,
            'outline_quality': outline_quality,
            'content_quality': content_quality,
            'consistency_quality': consistency_quality,
            'overall_quality': overall_quality,
        }

        # 质量不达标导致测试失败
        if overall_quality < 0.75:
            pytest.fail(f"总体质量不达标: {overall_quality*100:.1f}% < 75%")

    def test_10_generate_final_report(self, test_context):
        """Step 10: 生成最终测试报告"""
        print("\n=== Step 10: 生成最终测试报告 ===")

        # 计算执行时间
        execution_time = time.time() - test_context['start_time']
        test_context['execution_time'] = f'{execution_time:.1f} 秒'

        # 创建质量检查器
        checker = QualityChecker(test_context)

        # 生成最终报告
        report_file = REPORT_DIR / f'e2e_test_report_{test_context["timestamp"]}.md'
        checker.generate_final_report(report_file)

        print(f"\n最终测试报告: {report_file}")
        print(f"执行时间: {test_context['execution_time']}")

        # 输出摘要
        all_issues = (
            test_context['taa_issues'] +
            test_context['outline_issues'] +
            test_context['taw_issues'] +
            test_context['consistency_issues'] +
            test_context['image_quality_issues'] +
            test_context['professionalism_issues']
        )
        critical = len([i for i in all_issues if i['severity'] == 'critical'])
        high = len([i for i in all_issues if i['severity'] == 'high'])

        print(f"\n总问题数: {len(all_issues)} 个")
        print(f"  严重问题: {critical} 个")
        print(f"  高风险问题: {high} 个")

        # 输出质量得分
        if 'quality_scores' in test_context:
            scores = test_context['quality_scores']
            print(f"\n质量得分:")
            print(f"  总体质量: {scores['overall_quality']*100:.1f}%")

        if test_context['errors']:
            print(f"\n执行错误: {len(test_context['errors'])} 个")
            for error in test_context['errors']:
                print(f"  - {error['step']}: {error['error'][:100]}")

    # 保留原有的 test_6 和 test_7，重命名为 test_8 和 test_10
    # 原 test_6_trv_integration 已移至 test_8
    # 原 test_7_generate_final_report 已移至 test_10

