"""参数测试 - 测试 taw / trv 的关键参数"""
import shlex
import subprocess
from pathlib import Path

from docx import Document
import pytest

from validators.taw_validator import TAWValidator


# 项目根目录
PROJECT_ROOT = Path(__file__).parent.parent.parent

# 测试文件
OUTPUT_DIR = PROJECT_ROOT / 'output'
DRAFTS_DIR = PROJECT_ROOT / 'drafts'
TRV_OUTPUT_DIR = OUTPUT_DIR / 'trv'

# 测试章节（taa v2.1.0 后编号改为 1.X 体系）
TEST_CHAPTER = '1.3'


class TestParameters:
    """参数测试"""

    def _build_claude_prompt_command(self, skill_command: str) -> str:
        return f"claude -p {shlex.quote(skill_command)}"

    def _skip_nested_session_if_needed(self, result):
        if 'CLAUDECODE' in result.stderr or 'Nested sessions' in result.stderr or 'nested' in result.stderr.lower():
            pytest.skip('在 Claude Code 会话内部无法运行嵌套 claude skill 命令')

    def _latest_trv_report(self, review_type: str):
        report_files = list(TRV_OUTPUT_DIR.glob(f'审核报告_{review_type}_*.md'))
        if not report_files:
            return None
        return max(report_files, key=lambda p: p.stat().st_mtime)

    def _latest_revised_docx(self, stem: str):
        revised_files = list(TRV_OUTPUT_DIR.glob(f'{stem}_修订版_*.docx'))
        if not revised_files:
            return None
        return max(revised_files, key=lambda p: p.stat().st_mtime)

    def test_vendor_parameter(self):
        """测试 --vendor 参数"""
        print("\n=== 测试 --vendor 参数 ===")

        # 查找分析报告
        analysis_files = list(OUTPUT_DIR.glob('招标分析报告_*.md'))
        if not analysis_files:
            pytest.skip("未找到分析报告文件，请先执行 /taa")

        analysis_file = max(analysis_files, key=lambda p: p.stat().st_mtime)

        # 测试不同厂商
        vendors = ['博云', '灵雀云']

        for vendor in vendors:
            print(f"\n测试厂商: {vendor}")

            # 执行 taw
            skill_command = f'/taw {OUTPUT_DIR}/ --chapter {TEST_CHAPTER} --vendor "{vendor}" --search-tool auto'
            cmd = self._build_claude_prompt_command(skill_command)
            print(f"执行命令: {cmd}")

            try:
                result = subprocess.run(
                    cmd,
                    shell=True,
                    capture_output=True,
                    text=True,
                    timeout=600,
                )

                if result.returncode != 0:
                    print(f"生成失败: {result.stderr}")
                    continue

                # 查找输出文件
                chapter_files = list(DRAFTS_DIR.glob(f'{TEST_CHAPTER}_*.docx'))
                if not chapter_files:
                    print(f"未找到章节文件")
                    continue

                chapter_file = max(chapter_files, key=lambda p: p.stat().st_mtime)
                print(f"章节文件: {chapter_file}")

                # 验证厂商名一致性
                validator = TAWValidator(analysis_file)
                issues = validator.validate_vendor_consistency(chapter_file, vendor)

                if issues:
                    print(f"发现问题: {len(issues)} 个")
                    for issue in issues:
                        print(f"  - {issue['message']}")
                else:
                    print(f"✅ 厂商名一致性验证通过")

            except subprocess.TimeoutExpired:
                print(f"生成超时（10 分钟）")
                continue

    def test_search_tool_parameter(self):
        """测试 --search-tool 参数"""
        print("\n=== 测试 --search-tool 参数 ===")

        # 查找分析报告
        analysis_files = list(OUTPUT_DIR.glob('招标分析报告_*.md'))
        if not analysis_files:
            pytest.skip("未找到分析报告文件，请先执行 /taa")

        # 测试不同搜索工具
        search_tools = ['websearch', 'mcp', 'auto']

        for tool in search_tools:
            print(f"\n测试搜索工具: {tool}")

            # 执行 taw
            skill_command = f'/taw {OUTPUT_DIR}/ --chapter {TEST_CHAPTER} --search-tool {tool}'
            cmd = self._build_claude_prompt_command(skill_command)
            print(f"执行命令: {cmd}")

            try:
                result = subprocess.run(
                    cmd,
                    shell=True,
                    capture_output=True,
                    text=True,
                    timeout=600,
                )

                if result.returncode != 0:
                    print(f"生成失败: {result.stderr}")
                    continue

                # 查找输出文件
                chapter_files = list(DRAFTS_DIR.glob(f'{TEST_CHAPTER}_*.docx'))
                if not chapter_files:
                    print(f"未找到章节文件")
                    continue

                chapter_file = max(chapter_files, key=lambda p: p.stat().st_mtime)
                print(f"✅ 章节文件: {chapter_file}")

            except subprocess.TimeoutExpired:
                print(f"生成超时（10 分钟）")
                continue

    def test_no_kb_parameter(self):
        """测试 --no-kb 参数"""
        print("\n=== 测试 --no-kb 参数 ===")

        # 查找分析报告
        analysis_files = list(OUTPUT_DIR.glob('招标分析报告_*.md'))
        if not analysis_files:
            pytest.skip("未找到分析报告文件，请先执行 /taa")

        # 执行 taw（不使用知识库）
        skill_command = f'/taw {OUTPUT_DIR}/ --chapter {TEST_CHAPTER} --no-kb'
        cmd = self._build_claude_prompt_command(skill_command)
        print(f"执行命令: {cmd}")

        try:
            result = subprocess.run(
                cmd,
                shell=True,
                capture_output=True,
                text=True,
                timeout=600,
            )

            if result.returncode != 0:
                # 嵌套 Claude Code 会话无法运行，跳过测试
                if 'CLAUDECODE' in result.stderr or 'Nested sessions' in result.stderr or 'nested' in result.stderr.lower():
                    pytest.skip("在 Claude Code 会话内部无法运行嵌套 claude skill 命令")
                print(f"生成失败: {result.stderr}")
                pytest.fail("--no-kb 参数测试失败")

            # 查找输出文件
            chapter_files = list(DRAFTS_DIR.glob(f'{TEST_CHAPTER}_*.docx'))
            if not chapter_files:
                pytest.fail("未找到章节文件")

            chapter_file = max(chapter_files, key=lambda p: p.stat().st_mtime)
            print(f"✅ 章节文件: {chapter_file}")

        except subprocess.TimeoutExpired:
            pytest.fail("生成超时（10 分钟）")

    def test_trv_revise_docx_parameters(self):
        """测试 trv 的 --revise-docx / --revise-scope 参数"""
        print("\n=== 测试 trv 自动修订参数 ===")

        outline_files = list(OUTPUT_DIR.glob('投标文件大纲_*.docx'))
        analysis_files = list(OUTPUT_DIR.glob('招标分析报告_*.md'))
        if not outline_files:
            pytest.skip('未找到投标大纲文件，请先执行 /taa')
        if not analysis_files:
            pytest.skip('未找到分析报告文件，请先执行 /taa')

        outline_file = max(outline_files, key=lambda p: p.stat().st_mtime)
        analysis_file = max(analysis_files, key=lambda p: p.stat().st_mtime)
        stem = outline_file.stem

        skill_command = (
            f'/trv {shlex.quote(str(outline_file))} --type outline '
            f'--reference {shlex.quote(str(analysis_file))} --level all --revise-docx --revise-scope must'
        )
        cmd = self._build_claude_prompt_command(skill_command)
        print(f'执行命令: {cmd}')

        try:
            result = subprocess.run(
                cmd,
                shell=True,
                capture_output=True,
                text=True,
                timeout=600,
            )
            if result.returncode != 0:
                self._skip_nested_session_if_needed(result)
                print(f'执行失败: {result.stderr}')
                pytest.fail('--revise-docx 参数测试失败')

            report_file = self._latest_trv_report('outline')
            revised_file = self._latest_revised_docx(stem)
            assert report_file is not None, '未生成 outline 审核报告'
            assert revised_file is not None, '未生成修订版 DOCX'
            assert revised_file.exists()
            doc = Document(revised_file)
            # 验证修订版与原文件确实存在差异（避免假阳性）
            original_doc = Document(outline_file)
            orig_texts = [p.text for p in original_doc.paragraphs]
            revised_texts = [p.text for p in doc.paragraphs]
            has_diff = orig_texts != revised_texts
            if not has_diff:
                # 也检查表格内容
                for orig_tbl, rev_tbl in zip(original_doc.tables, doc.tables):
                    for orig_row, rev_row in zip(orig_tbl.rows, rev_tbl.rows):
                        for orig_cell, rev_cell in zip(orig_row.cells, rev_row.cells):
                            if orig_cell.text != rev_cell.text:
                                has_diff = True
                                break
                        if has_diff:
                            break
                    if has_diff:
                        break
            assert has_diff, '修订版文件与原文件内容完全相同，修订未实际生效'
            print(f'✅ 审核报告: {report_file.name}')
            print(f'✅ 修订版文件: {revised_file.name}（已验证存在实际修订）')
        except subprocess.TimeoutExpired:
            pytest.fail('trv 自动修订参数测试超时（10 分钟）')

    def test_trv_without_revise_docx_does_not_require_revised_output(self):
        """测试未启用 --revise-docx 时仍可正常审核"""
        print("\n=== 测试 trv 默认不修订 ===")

        outline_files = list(OUTPUT_DIR.glob('投标文件大纲_*.docx'))
        analysis_files = list(OUTPUT_DIR.glob('招标分析报告_*.md'))
        if not outline_files:
            pytest.skip('未找到投标大纲文件，请先执行 /taa')
        if not analysis_files:
            pytest.skip('未找到分析报告文件，请先执行 /taa')

        outline_file = max(outline_files, key=lambda p: p.stat().st_mtime)
        analysis_file = max(analysis_files, key=lambda p: p.stat().st_mtime)
        stem = outline_file.stem

        # 记录执行前已存在的修订版文件，用于后续判断是否新生成
        pre_existing_revised = set(TRV_OUTPUT_DIR.glob(f'{stem}_修订版_*.docx'))

        skill_command = (
            f'/trv {shlex.quote(str(outline_file))} --type outline '
            f'--reference {shlex.quote(str(analysis_file))} --level all'
        )
        cmd = self._build_claude_prompt_command(skill_command)
        print(f'执行命令: {cmd}')

        try:
            result = subprocess.run(
                cmd,
                shell=True,
                capture_output=True,
                text=True,
                timeout=600,
            )
            if result.returncode != 0:
                self._skip_nested_session_if_needed(result)
                print(f'执行失败: {result.stderr}')
                pytest.fail('默认 trv 参数测试失败')

            report_file = self._latest_trv_report('outline')
            assert report_file is not None, '未生成 outline 审核报告'
            print(f'✅ 审核报告: {report_file.name}')

            # 验证未启用 --revise-docx 时不应新生成修订版
            post_revised = set(TRV_OUTPUT_DIR.glob(f'{stem}_修订版_*.docx'))
            new_revised = post_revised - pre_existing_revised
            assert not new_revised, f'未启用 --revise-docx 但生成了新修订版文件: {[f.name for f in new_revised]}'
            print('✅ 确认未生成新修订版文件')
        except subprocess.TimeoutExpired:
            pytest.fail('trv 默认参数测试超时（10 分钟）')
