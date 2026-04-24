"""DOCX 文件解析工具"""
import re
from pathlib import Path
from typing import Dict, List, Any
from docx import Document
from docx.oxml.ns import qn


class DOCXParser:
    """DOCX 文件解析器"""

    @staticmethod
    def parse_document(docx_file: Path) -> Dict[str, Any]:
        """解析 DOCX 文档结构"""
        doc = Document(docx_file)
        return {
            'file': str(docx_file),
            'text': DOCXParser.extract_text(doc),
            'word_count': DOCXParser.count_words(DOCXParser.extract_text(doc)),
            'images': DOCXParser.extract_images(doc),
            'tables': DOCXParser.extract_tables(doc),
            'paragraphs': [p.text for p in doc.paragraphs if p.text.strip()],
        }

    @staticmethod
    def extract_text(doc: Document) -> str:
        """提取纯文本"""
        return '\n'.join([p.text for p in doc.paragraphs])

    @staticmethod
    def count_words(text: str) -> int:
        """统计字数（中文字符 + 英文单词）"""
        # 统计中文字符
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        # 统计英文单词
        english_words = len(re.findall(r'\b[a-zA-Z]+\b', text))
        return chinese_chars + english_words

    @staticmethod
    def extract_images(doc: Document) -> List[Dict[str, Any]]:
        """提取图片信息"""
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                images.append({
                    'type': 'image',
                    'target': rel.target_ref,
                })

        # 统计 InlineShape（嵌入图片）
        for para in doc.paragraphs:
            for run in para.runs:
                if run._element.xpath('.//a:blip'):
                    images.append({
                        'type': 'inline_shape',
                        'paragraph': para.text[:50],
                    })

        return images

    @staticmethod
    def extract_tables(doc: Document) -> List[Dict[str, Any]]:
        """提取表格"""
        tables = []
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            tables.append({
                'index': i,
                'rows': len(table.rows),
                'cols': len(table.columns),
                'data': rows,
            })
        return tables

    @staticmethod
    def parse_outline_structure(doc: Document) -> List[Dict[str, Any]]:
        """解析大纲结构（识别标题层级）"""
        structure = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name

            # 识别标题层级
            if style_name.startswith('Heading'):
                level = int(style_name.replace('Heading', '').strip() or '1')
                structure.append({
                    'level': level,
                    'text': text,
                    'style': style_name,
                })
            # 识别中文编号
            elif re.match(r'^[一二三四五六七八九十]+、', text):
                structure.append({
                    'level': 1,
                    'text': text,
                    'style': 'chinese_number',
                })
            elif re.match(r'^\d+\.\d+', text):
                level = text.count('.') + 1
                structure.append({
                    'level': level,
                    'text': text,
                    'style': 'decimal_number',
                })

        return structure
