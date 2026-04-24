#!/usr/bin/env python3
"""
tests/test_kb_index.py

知识库索引功能测试（4项）：
  Test 1: DOCX 文本提取管道（生成 fake DOCX → 提取文字）
  Test 2: PDF 文本提取管道（生成 fake PDF → pdfplumber 提取文字）
  Test 3: kb_indexer.py --help 不报错
  Test 4: kb_indexer.py --scan 对空目录不报错
"""

import os
import sys
import subprocess
import tempfile

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
KB_INDEXER = os.path.join(PROJECT_ROOT, 'skills', 'taw', 'tools', 'kb_indexer.py')
sys.path.insert(0, PROJECT_ROOT)


def ensure_deps():
    pkgs = ['python-docx', 'pdfplumber', 'pyyaml']
    for pkg in pkgs:
        try:
            if pkg == 'python-docx':
                import docx  # noqa: F401
            elif pkg == 'pdfplumber':
                import pdfplumber  # noqa: F401
            elif pkg == 'pyyaml':
                import yaml  # noqa: F401
        except ImportError:
            subprocess.run([sys.executable, '-m', 'pip', 'install', pkg, '-q',
                            '--break-system-packages'], check=False)


def test_1_docx_text_extraction():
    """Test 1: DOCX 文本提取管道。"""
    ensure_deps()
    from docx import Document

    with tempfile.TemporaryDirectory() as tmpdir:
        # 创建包含已知文本的 fake DOCX
        docx_path = os.path.join(tmpdir, 'fake.docx')
        doc = Document()
        doc.add_heading('测试标题', level=1)
        doc.add_paragraph('这是测试段落内容，包含技术关键词：Kubernetes、容器云、微服务。')
        doc.save(docx_path)

        # 验证文件创建成功
        assert os.path.exists(docx_path), f'DOCX 文件未创建：{docx_path}'
        assert os.path.getsize(docx_path) > 0, 'DOCX 文件为空'

        # 重新读取验证内容
        doc2 = Document(docx_path)
        all_text = '\n'.join([p.text for p in doc2.paragraphs])

    assert 'Kubernetes' in all_text, f'提取文本应含 "Kubernetes"，实际：{all_text[:200]!r}'
    assert '容器云' in all_text, f'提取文本应含 "容器云"，实际：{all_text[:200]!r}'
    print(f'Test 1 PASSED: DOCX 文本提取正常（{len(all_text)} 字）')


def test_2_pdf_text_extraction():
    """Test 2: PDF 文本提取管道（使用 fpdf2 生成 fake PDF）。"""
    ensure_deps()

    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_path = os.path.join(tmpdir, 'fake.pdf')

        try:
            from fpdf import FPDF
        except ImportError:
            subprocess.run([sys.executable, '-m', 'pip', 'install', 'fpdf2', '-q',
                            '--break-system-packages'], check=False)
            try:
                from fpdf import FPDF
            except ImportError:
                print('Test 2 SKIPPED: fpdf2 不可用，跳过 PDF 文本提取测试')
                return

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font('Helvetica', size=12)
        pdf.cell(200, 10, txt='Test PDF Content: Kubernetes container cloud', ln=True)
        pdf.output(pdf_path)

        import pdfplumber
        with pdfplumber.open(pdf_path) as pdf_reader:
            text = '\n'.join([page.extract_text() or '' for page in pdf_reader.pages])

    assert 'Kubernetes' in text or 'Test' in text, \
        f'PDF 文本提取应包含测试内容，实际：{text[:200]!r}'
    print(f'Test 2 PASSED: PDF 文本提取正常（{len(text)} 字）')


def test_3_kb_indexer_help():
    """Test 3: kb_indexer.py --help 正常退出（returncode=0）。"""
    if not os.path.exists(KB_INDEXER):
        print('Test 3 SKIPPED: kb_indexer.py 不存在')
        return

    result = subprocess.run(
        [sys.executable, KB_INDEXER, '--help'],
        capture_output=True, text=True
    )
    assert result.returncode == 0, \
        f'--help 返回码应为 0，实际：{result.returncode}\nstderr: {result.stderr}'
    print('Test 3 PASSED: kb_indexer.py --help 正常')


def test_4_kb_indexer_empty_dir():
    """Test 4: kb_indexer.py --scan 对空目录正常退出（无文档时 exit 1）。"""
    if not os.path.exists(KB_INDEXER):
        print('Test 4 SKIPPED: kb_indexer.py 不存在')
        return

    with tempfile.TemporaryDirectory() as tmpdir:
        result = subprocess.run(
            [sys.executable, KB_INDEXER, '--scan', '--kb-path', tmpdir],
            capture_output=True, text=True
        )
        # 空目录无文档，预期 exit 1 并输出提示
        assert result.returncode == 1, \
            f'空目录扫描应返回 1，实际：{result.returncode}'
        assert '未找到' in result.stdout, \
            f'应输出"未找到"提示，实际：{result.stdout}'
    print('Test 4 PASSED: kb_indexer.py 空目录扫描正常退出')


if __name__ == '__main__':
    ensure_deps()
    tests = [
        test_1_docx_text_extraction,
        test_2_pdf_text_extraction,
        test_3_kb_indexer_help,
        test_4_kb_indexer_empty_dir,
    ]

    failures = []
    for test in tests:
        try:
            test()
        except Exception as e:
            failures.append((test.__name__, e))
            print(f'{test.__name__} FAILED: {e}')

    print()
    if failures:
        print(f'{len(failures)}/{len(tests)} 个测试失败')
        sys.exit(1)
    else:
        print(f'全部 {len(tests)} 个测试通过')
