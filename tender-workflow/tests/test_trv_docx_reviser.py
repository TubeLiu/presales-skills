"""tests/test_trv_docx_reviser.py — 指令驱动 DOCX 修订工具单元测试"""
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from tools.trv_docx_reviser import revise_docx, apply_instructions, ReviseStats


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _write_instructions(path: Path, instructions: list, scope: str = 'must') -> None:
    data = {
        'meta': {'scope': scope, 'instruction_count': len(instructions)},
        'instructions': instructions,
    }
    path.write_text(json.dumps(data, ensure_ascii=False), encoding='utf-8')


def _build_simple_docx(path: Path) -> None:
    """构造含段落和表格的简单 DOCX 测试文件。"""
    doc = Document()
    doc.add_paragraph('本次项目交付结束后，提供完善的售后服务。')
    doc.add_paragraph('ARM 的 CPU 架构存在部分功能不兼容。')
    p3 = doc.add_paragraph('')
    run1 = p3.add_run('前半段文本')
    run1.bold = True
    run2 = p3.add_run('需要替换的内容')
    run2.italic = True
    run3 = p3.add_run('后半段文本')
    run3.font.color.rgb = RGBColor(0xFF, 0, 0)

    table = doc.add_table(rows=3, cols=4)
    headers = ['序号', '技术内容', '投标规格', '偏离']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].text = '1'
    table.rows[1].cells[1].text = '平台兼容性'
    table.rows[1].cells[2].text = '完全满足'
    table.rows[1].cells[3].text = '无偏离'
    table.rows[2].cells[0].text = '2'
    table.rows[2].cells[1].text = '售后服务'
    table.rows[2].cells[2].text = '标准维保'
    table.rows[2].cells[3].text = '无偏离'
    doc.save(path)


# ---------------------------------------------------------------------------
# Tests: paragraph_text_replace
# ---------------------------------------------------------------------------

def test_paragraph_text_replace(tmp_path: Path):
    """段落内查找替换，保留格式。"""
    input_path = tmp_path / 'input.docx'
    instructions_path = tmp_path / 'instructions.json'
    _build_simple_docx(input_path)
    _write_instructions(instructions_path, [
        {
            'id': 'R001',
            'severity': 'critical',
            'type': 'paragraph_text_replace',
            'description': '修正售后服务描述',
            'match': {'method': 'contains', 'text': '售后服务'},
            'action': {
                'find': '提供完善的售后服务',
                'replace': '提供 1 年 7*24 小时原厂技术支持服务',
            },
        },
    ])

    result = revise_docx(input_path, instructions_path, tmp_path / 'output')
    output_path = Path(result['output_file'])
    assert output_path.exists()
    assert result['summary']['applied'] == 1
    assert result['summary']['not_found'] == 0

    doc = Document(output_path)
    assert '1 年 7*24 小时原厂技术支持服务' in doc.paragraphs[0].text


def test_paragraph_text_replace_not_found(tmp_path: Path):
    """查找文本不存在时返回 not_found。"""
    input_path = tmp_path / 'input.docx'
    instructions_path = tmp_path / 'instructions.json'
    _build_simple_docx(input_path)
    _write_instructions(instructions_path, [
        {
            'id': 'R001',
            'severity': 'critical',
            'type': 'paragraph_text_replace',
            'description': '替换不存在的文本',
            'match': {'method': 'contains', 'text': '完全不存在的内容'},
            'action': {'find': '完全不存在的内容', 'replace': '新内容'},
        },
    ])

    result = revise_docx(input_path, instructions_path, tmp_path / 'output')
    assert result['summary']['not_found'] == 1
    assert result['summary']['applied'] == 0


# ---------------------------------------------------------------------------
# Tests: run-level format preservation
# ---------------------------------------------------------------------------

def test_replace_preserves_run_formatting(tmp_path: Path):
    """跨 run 替换时，保留首 run 的格式属性。"""
    input_path = tmp_path / 'input.docx'
    instructions_path = tmp_path / 'instructions.json'
    _build_simple_docx(input_path)
    _write_instructions(instructions_path, [
        {
            'id': 'R001',
            'severity': 'major',
            'type': 'paragraph_text_replace',
            'description': '替换跨 run 文本',
            'match': {'method': 'contains', 'text': '需要替换的内容'},
            'action': {
                'find': '需要替换的内容',
                'replace': '已替换内容',
            },
        },
    ])

    result = revise_docx(input_path, instructions_path, tmp_path / 'output')
    assert result['summary']['applied'] == 1

    doc = Document(Path(result['output_file']))
    # 第三段（索引 2）包含多 run
    para = doc.paragraphs[2]
    full_text = ''.join(r.text for r in para.runs)
    assert '已替换内容' in full_text
    # 首 run 应保留 bold 格式
    assert para.runs[0].bold is True


# ---------------------------------------------------------------------------
# Tests: paragraph_full_replace
# ---------------------------------------------------------------------------

def test_paragraph_full_replace(tmp_path: Path):
    """整段替换。"""
    input_path = tmp_path / 'input.docx'
    instructions_path = tmp_path / 'instructions.json'
    _build_simple_docx(input_path)
    _write_instructions(instructions_path, [
        {
            'id': 'R001',
            'severity': 'critical',
            'type': 'paragraph_full_replace',
            'description': '整段替换',
            'match': {'method': 'contains', 'text': 'ARM'},
            'action': {'new_text': 'ARM 架构兼容性已完成验证，满足交付要求。'},
        },
    ])

    result = revise_docx(input_path, instructions_path, tmp_path / 'output')
    assert result['summary']['applied'] == 1

    doc = Document(Path(result['output_file']))
    assert doc.paragraphs[1].text == 'ARM 架构兼容性已完成验证，满足交付要求。'


# ---------------------------------------------------------------------------
# Tests: table_cell_replace
# ---------------------------------------------------------------------------

def test_table_cell_replace(tmp_path: Path):
    """按表头+行号定位单元格替换。"""
    input_path = tmp_path / 'input.docx'
    instructions_path = tmp_path / 'instructions.json'
    _build_simple_docx(input_path)
    _write_instructions(instructions_path, [
        {
            'id': 'R001',
            'severity': 'critical',
            'type': 'table_cell_replace',
            'description': '修正偏离表',
            'match': {
                'method': 'table_header_row',
                'header_contains': ['序号', '投标规格'],
                'row_match': {'column': '序号', 'value': '1'},
                'target_column': '投标规格',
            },
            'action': {
                'find': '完全满足',
                'replace': '满足，已完成兼容性验证',
            },
        },
    ])

    result = revise_docx(input_path, instructions_path, tmp_path / 'output')
    assert result['summary']['applied'] == 1

    doc = Document(Path(result['output_file']))
    cell_text = doc.tables[0].rows[1].cells[2].text
    assert '已完成兼容性验证' in cell_text


def test_table_cell_not_found(tmp_path: Path):
    """表格单元格不存在时返回 not_found。"""
    input_path = tmp_path / 'input.docx'
    instructions_path = tmp_path / 'instructions.json'
    _build_simple_docx(input_path)
    _write_instructions(instructions_path, [
        {
            'id': 'R001',
            'severity': 'critical',
            'type': 'table_cell_replace',
            'description': '查找不存在的表格',
            'match': {
                'method': 'table_header_row',
                'header_contains': ['不存在的表头'],
                'row_match': {'column': '序号', 'value': '99'},
                'target_column': '不存在',
            },
            'action': {'find': 'x', 'replace': 'y'},
        },
    ])

    result = revise_docx(input_path, instructions_path, tmp_path / 'output')
    assert result['summary']['not_found'] == 1


# ---------------------------------------------------------------------------
# Tests: global_text_replace
# ---------------------------------------------------------------------------

def test_global_text_replace(tmp_path: Path):
    """全文查找替换。"""
    # 构造含特定文本的 DOCX
    input_path = tmp_path / 'input.docx'
    doc = Document()
    doc.add_paragraph('贵行项目管理办法要求贵行确认。')
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = '项目'
    table.rows[0].cells[1].text = '说明'
    table.rows[1].cells[0].text = '平台'
    table.rows[1].cells[1].text = '贵行指定的平台。'
    doc.save(input_path)

    instructions_path = tmp_path / 'instructions.json'
    _write_instructions(instructions_path, [
        {
            'id': 'R001',
            'severity': 'major',
            'type': 'global_text_replace',
            'description': '统一术语',
            'match': {'method': 'global'},
            'action': {'find': '贵行', 'replace': '招标方'},
        },
    ])

    result = revise_docx(input_path, instructions_path, tmp_path / 'output')
    assert result['summary']['applied'] == 1

    doc = Document(Path(result['output_file']))
    assert '贵行' not in doc.paragraphs[0].text
    assert '招标方' in doc.paragraphs[0].text
    assert '贵行' not in doc.tables[0].rows[1].cells[1].text


# ---------------------------------------------------------------------------
# Tests: paragraph_insert_after
# ---------------------------------------------------------------------------

def test_paragraph_insert_after(tmp_path: Path):
    """在匹配段落后插入新段落。"""
    input_path = tmp_path / 'input.docx'
    instructions_path = tmp_path / 'instructions.json'
    _build_simple_docx(input_path)
    _write_instructions(instructions_path, [
        {
            'id': 'R001',
            'severity': 'major',
            'type': 'paragraph_insert_after',
            'description': '补充说明',
            'match': {'method': 'contains', 'text': '售后服务'},
            'action': {'insert_text': '补充：含 7*24 小时技术支持。'},
        },
    ])

    result = revise_docx(input_path, instructions_path, tmp_path / 'output')
    assert result['summary']['applied'] == 1

    doc = Document(Path(result['output_file']))
    texts = [p.text for p in doc.paragraphs]
    assert '补充：含 7*24 小时技术支持。' in texts


# ---------------------------------------------------------------------------
# Tests: paragraph_delete
# ---------------------------------------------------------------------------

def test_paragraph_delete(tmp_path: Path):
    """删除匹配段落。"""
    input_path = tmp_path / 'input.docx'
    instructions_path = tmp_path / 'instructions.json'
    _build_simple_docx(input_path)

    # 确认删除前存在
    doc = Document(input_path)
    assert any('ARM' in p.text for p in doc.paragraphs)

    _write_instructions(instructions_path, [
        {
            'id': 'R001',
            'severity': 'minor',
            'type': 'paragraph_delete',
            'description': '删除不适用段落',
            'match': {'method': 'contains', 'text': 'ARM'},
        },
    ])

    result = revise_docx(input_path, instructions_path, tmp_path / 'output')
    assert result['summary']['applied'] == 1

    doc = Document(Path(result['output_file']))
    assert not any('ARM' in p.text for p in doc.paragraphs)


# ---------------------------------------------------------------------------
# Tests: mixed instructions
# ---------------------------------------------------------------------------

def test_multiple_instructions(tmp_path: Path):
    """多条混合指令按顺序执行。"""
    input_path = tmp_path / 'input.docx'
    instructions_path = tmp_path / 'instructions.json'
    _build_simple_docx(input_path)
    _write_instructions(instructions_path, [
        {
            'id': 'R001',
            'severity': 'critical',
            'type': 'paragraph_text_replace',
            'description': '修正售后',
            'match': {'method': 'contains', 'text': '售后服务'},
            'action': {'find': '完善的售后服务', 'replace': '原厂技术支持'},
        },
        {
            'id': 'R002',
            'severity': 'critical',
            'type': 'table_cell_replace',
            'description': '修正偏离表',
            'match': {
                'method': 'table_header_row',
                'header_contains': ['序号', '投标规格'],
                'row_match': {'column': '序号', 'value': '2'},
                'target_column': '投标规格',
            },
            'action': {'find': '标准维保', 'replace': '原厂 7*24 服务'},
        },
        {
            'id': 'R003',
            'severity': 'minor',
            'type': 'paragraph_text_replace',
            'description': '不存在的替换',
            'match': {'method': 'contains', 'text': '完全不存在'},
            'action': {'find': '完全不存在', 'replace': 'xxx'},
        },
    ])

    result = revise_docx(input_path, instructions_path, tmp_path / 'output')
    assert result['summary']['total'] == 3
    assert result['summary']['applied'] == 2
    assert result['summary']['not_found'] == 1


# ---------------------------------------------------------------------------
# Tests: unknown type
# ---------------------------------------------------------------------------

def test_unknown_instruction_type(tmp_path: Path):
    """未知指令类型返回 error。"""
    input_path = tmp_path / 'input.docx'
    instructions_path = tmp_path / 'instructions.json'
    _build_simple_docx(input_path)
    _write_instructions(instructions_path, [
        {
            'id': 'R001',
            'severity': 'critical',
            'type': 'nonexistent_type',
            'description': '未知类型',
            'match': {},
            'action': {},
        },
    ])

    result = revise_docx(input_path, instructions_path, tmp_path / 'output')
    assert result['summary']['errors'] == 1


# ---------------------------------------------------------------------------
# Tests: empty instructions
# ---------------------------------------------------------------------------

def test_empty_instructions(tmp_path: Path):
    """空指令列表仍生成修订版文件（内容不变）。"""
    input_path = tmp_path / 'input.docx'
    instructions_path = tmp_path / 'instructions.json'
    _build_simple_docx(input_path)
    _write_instructions(instructions_path, [])

    result = revise_docx(input_path, instructions_path, tmp_path / 'output')
    assert Path(result['output_file']).exists()
    assert result['summary']['total'] == 0


# ---------------------------------------------------------------------------
# Tests: regex matching
# ---------------------------------------------------------------------------

def test_regex_matching(tmp_path: Path):
    """正则表达式匹配段落。"""
    input_path = tmp_path / 'input.docx'
    instructions_path = tmp_path / 'instructions.json'
    _build_simple_docx(input_path)
    _write_instructions(instructions_path, [
        {
            'id': 'R001',
            'severity': 'major',
            'type': 'paragraph_full_replace',
            'description': '正则匹配测试',
            'match': {'method': 'regex', 'pattern': r'ARM.*不兼容'},
            'action': {'new_text': 'ARM 架构已通过兼容性验证。'},
        },
    ])

    result = revise_docx(input_path, instructions_path, tmp_path / 'output')
    assert result['summary']['applied'] == 1

    doc = Document(Path(result['output_file']))
    assert doc.paragraphs[1].text == 'ARM 架构已通过兼容性验证。'


# ---------------------------------------------------------------------------
# Tests: global replace where find is substring of replace
# ---------------------------------------------------------------------------

def test_context_before_disambiguation(tmp_path: Path):
    """context_before 用于同段落消歧。"""
    input_path = tmp_path / 'input.docx'
    doc = Document()
    doc.add_paragraph('项目交付结束后，提供完善的售后服务。')
    doc.add_paragraph('培训结束后，提供完善的售后服务。')
    doc.save(input_path)

    instructions_path = tmp_path / 'instructions.json'
    _write_instructions(instructions_path, [
        {
            'id': 'R001',
            'severity': 'critical',
            'type': 'paragraph_text_replace',
            'description': '仅修改包含"培训"的段落',
            'match': {
                'method': 'contains',
                'text': '售后服务',
                'context_before': '培训',
            },
            'action': {'find': '完善的售后服务', 'replace': '原厂技术支持'},
        },
    ])

    result = revise_docx(input_path, instructions_path, tmp_path / 'output')
    assert result['summary']['applied'] == 1

    doc = Document(Path(result['output_file']))
    # 第一段不应被修改
    assert '完善的售后服务' in doc.paragraphs[0].text
    # 第二段应被修改
    assert '原厂技术支持' in doc.paragraphs[1].text


def test_exact_matching(tmp_path: Path):
    """exact 方法精确匹配段落。"""
    input_path = tmp_path / 'input.docx'
    doc = Document()
    doc.add_paragraph('售后服务方案')
    doc.add_paragraph('售后服务方案详述如下')
    doc.save(input_path)

    instructions_path = tmp_path / 'instructions.json'
    _write_instructions(instructions_path, [
        {
            'id': 'R001',
            'severity': 'major',
            'type': 'paragraph_full_replace',
            'description': '精确匹配短段落',
            'match': {'method': 'exact', 'text': '售后服务方案'},
            'action': {'new_text': '技术支持方案'},
        },
    ])

    result = revise_docx(input_path, instructions_path, tmp_path / 'output')
    assert result['summary']['applied'] == 1

    doc = Document(Path(result['output_file']))
    assert doc.paragraphs[0].text == '技术支持方案'
    # 第二段不应被匹配（包含额外文本）
    assert '详述如下' in doc.paragraphs[1].text


def test_table_cell_full_replacement(tmp_path: Path):
    """无 find_text 时直接替换整个单元格。"""
    input_path = tmp_path / 'input.docx'
    instructions_path = tmp_path / 'instructions.json'
    _build_simple_docx(input_path)
    _write_instructions(instructions_path, [
        {
            'id': 'R001',
            'severity': 'critical',
            'type': 'table_cell_replace',
            'description': '整单元格替换',
            'match': {
                'method': 'table_header_row',
                'header_contains': ['序号', '投标规格'],
                'row_match': {'column': '序号', 'value': '1'},
                'target_column': '投标规格',
            },
            'action': {'replace': '满足，全面兼容'},
        },
    ])

    result = revise_docx(input_path, instructions_path, tmp_path / 'output')
    assert result['summary']['applied'] == 1

    doc = Document(Path(result['output_file']))
    assert doc.tables[0].rows[1].cells[2].text == '满足，全面兼容'


def test_cross_three_runs_replacement(tmp_path: Path):
    """跨 3 个 run 的替换。"""
    input_path = tmp_path / 'input.docx'
    doc = Document()
    p = doc.add_paragraph('')
    r1 = p.add_run('开始')
    r1.bold = True
    r2 = p.add_run('中间需要')
    r2.italic = True
    r3 = p.add_run('替换的文本')
    r3.font.size = Pt(14)
    r4 = p.add_run('结尾')
    r4.font.color.rgb = RGBColor(0, 0, 0xFF)
    doc.save(input_path)

    instructions_path = tmp_path / 'instructions.json'
    _write_instructions(instructions_path, [
        {
            'id': 'R001',
            'severity': 'major',
            'type': 'paragraph_text_replace',
            'description': '跨 3 run 替换',
            'match': {'method': 'contains', 'text': '中间'},
            'action': {'find': '中间需要替换的文本', 'replace': '新内容'},
        },
    ])

    result = revise_docx(input_path, instructions_path, tmp_path / 'output')
    assert result['summary']['applied'] == 1

    doc = Document(Path(result['output_file']))
    p = doc.paragraphs[0]
    full = ''.join(r.text for r in p.runs)
    assert full == '开始新内容结尾'
    # 首 run 保留 bold
    assert p.runs[0].bold is True


def test_global_replace_find_in_replace(tmp_path: Path):
    """find_text 是 replace_text 的子串时不应无限循环。"""
    input_path = tmp_path / 'input.docx'
    doc = Document()
    doc.add_paragraph('版本 A 已发布。')
    doc.add_paragraph('版本 A 需要更新。')
    doc.save(input_path)

    instructions_path = tmp_path / 'instructions.json'
    _write_instructions(instructions_path, [
        {
            'id': 'R001',
            'severity': 'major',
            'type': 'global_text_replace',
            'description': '扩展版本号',
            'match': {'method': 'global'},
            'action': {'find': '版本 A', 'replace': '版本 A v2.0'},
        },
    ])

    result = revise_docx(input_path, instructions_path, tmp_path / 'output')
    assert result['summary']['applied'] == 1

    doc = Document(Path(result['output_file']))
    assert '版本 A v2.0 已发布' in doc.paragraphs[0].text
    # 确认没有被重复替换
    assert doc.paragraphs[0].text.count('v2.0') == 1
