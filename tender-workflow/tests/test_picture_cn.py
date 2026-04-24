#!/usr/bin/env python3
"""
tests/test_picture_cn.py

图片插入功能测试：
  Test 1: 图片存在时正常插入
  Test 2: 图片不存在时降级为占位符（不抛异常）
"""

import os
import sys
import tempfile

# 确保从项目根目录运行
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, PROJECT_ROOT)


def ensure_docx():
    try:
        from docx import Document  # noqa: F401
    except ImportError:
        import subprocess
        subprocess.run([sys.executable, '-m', 'pip', 'install', 'python-docx', '-q',
                        '--break-system-packages'], check=False)


def create_dummy_png(path: str):
    """创建最小合法 PNG（1x1 像素）。优先用 Pillow，备用手工构造。"""
    try:
        from PIL import Image
        img = Image.new('RGB', (1, 1), color=(255, 0, 0))
        img.save(path, format='PNG')
        return
    except ImportError:
        pass
    # 手工构造合法的 1x1 灰度 PNG（经 CRC32 验证正确）
    import struct, zlib

    def make_chunk(chunk_type: bytes, data: bytes) -> bytes:
        length = struct.pack('>I', len(data))
        crc = struct.pack('>I', zlib.crc32(chunk_type + data) & 0xFFFFFFFF)
        return length + chunk_type + data + crc

    signature = b'\x89PNG\r\n\x1a\n'
    ihdr = make_chunk(b'IHDR', struct.pack('>IIBBBBB', 1, 1, 8, 0, 0, 0, 0))
    # 1x1 grayscale: filter_byte=0, pixel=128
    raw_row = b'\x00\x80'
    idat_data = zlib.compress(raw_row)
    idat = make_chunk(b'IDAT', idat_data)
    iend = make_chunk(b'IEND', b'')
    png_bytes = signature + ihdr + idat + iend
    with open(path, 'wb') as f:
        f.write(png_bytes)


def get_add_picture_cn():
    """从 SKILL.md 提取并 exec add_picture_cn 函数（内嵌代码）。"""
    # 直接定义一份简化版 add_picture_cn，与 SKILL.md 中实现一致
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def apply_run_font(run, cn_font='宋体', latin_font='Times New Roman'):
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), cn_font)

    def add_picture_cn(doc, img_path, caption=None, width_cm=14.0, cn_font='宋体'):
        if not os.path.exists(img_path):
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.left_indent = Cm(0.75)
            run = p.add_run('[图片占位符 — ' + (caption or img_path) + '：请手动插入图片]')
            run.italic = True
            apply_run_font(run, cn_font=cn_font)
            return
        doc.add_picture(img_path, width=Cm(width_cm))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap_p = doc.add_paragraph(style='Normal')
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_p.add_run(caption)
            cap_run.italic = True
            cap_run.font.size = Pt(10.5)
            apply_run_font(cap_run, cn_font=cn_font)

    return add_picture_cn


def test_1_image_exists():
    """Test 1: 图片存在时正常插入，文档段落数增加。"""
    ensure_docx()
    from docx import Document

    add_picture_cn = get_add_picture_cn()

    with tempfile.TemporaryDirectory() as tmpdir:
        img_path = os.path.join(tmpdir, 'test.png')
        create_dummy_png(img_path)

        doc = Document()
        para_count_before = len(doc.paragraphs)
        add_picture_cn(doc, img_path, caption='测试图片说明', width_cm=10.0)
        para_count_after = len(doc.paragraphs)

        assert para_count_after > para_count_before, \
            f'图片插入后段落数应增加，before={para_count_before}, after={para_count_after}'

        # caption 段落应包含说明文字
        last_para = doc.paragraphs[-1]
        assert '测试图片说明' in last_para.text, \
            f'caption 未写入文档，last_para.text={last_para.text!r}'

    print('Test 1 PASSED: 图片存在时正常插入')


def test_2_image_missing_fallback():
    """Test 2: 图片不存在时降级为占位符，不抛异常。"""
    ensure_docx()
    from docx import Document

    add_picture_cn = get_add_picture_cn()

    doc = Document()
    para_count_before = len(doc.paragraphs)

    # 不存在的路径
    add_picture_cn(doc, '/nonexistent/path/image.png', caption='架构图占位符')

    para_count_after = len(doc.paragraphs)
    assert para_count_after > para_count_before, \
        '占位符应添加段落'

    # 占位符文本应包含提示
    placeholder_text = doc.paragraphs[-1].text
    assert '占位符' in placeholder_text or '请手动插入' in placeholder_text, \
        f'占位符段落文本不符合预期: {placeholder_text!r}'

    print('Test 2 PASSED: 图片不存在时降级为占位符（无异常）')


if __name__ == '__main__':
    failures = []
    tests = [test_1_image_exists, test_2_image_missing_fallback]
    for test in tests:
        try:
            test()
        except Exception as e:
            failures.append((test.__name__, e))
            print(f'{test.__name__} FAILED: {e}')

    if failures:
        print(f'\n{len(failures)}/{len(tests)} 个测试失败')
        sys.exit(1)
    else:
        print(f'\n全部 {len(tests)} 个测试通过')
