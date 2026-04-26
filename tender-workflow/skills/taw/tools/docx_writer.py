#!/usr/bin/env python3
"""
taw DOCX 生成工具库

提供中文排版的 DOCX 生成函数，替代在 SKILL.md 中内嵌代码模板。
所有正文写入统一使用 write_markdown()，禁止直接使用 doc.add_heading() / doc.add_paragraph()。

用法（推荐：脚本内部用 Path(__file__) 自定位，调用方只传脚本绝对路径）：
    # 在 SKILL.md 中：先按 §路径自定位 解析 SKILL_DIR，然后：
    #   python3 "$SKILL_DIR/tools/docx_writer.py" ...  （直接当 CLI 跑）
    # 或在其它 Python 中 import：
    #   import sys; from pathlib import Path
    #   sys.path.insert(0, str(Path("$SKILL_DIR/tools")))
    from docx_writer import create_document, write_markdown, add_heading_cn, add_picture_cn

源码模式直接从 tender-workflow/ 根目录运行时：
    import sys; sys.path.insert(0, 'skills/taw/tools')
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 样式级字体设置 ──────────────────────────────────

def _apply_font(style, size_pt, bold=False, cn_font='宋体', latin_font='Times New Roman'):
    """设置样式的字体、字号、加粗"""
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.font.name = latin_font
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), latin_font)
    rFonts.set(qn('w:hAnsi'), latin_font)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:cs'), cn_font)
    # 删除主题字体引用（优先级高于明确字体名，必须清除）
    for attr in [qn('w:asciiTheme'), qn('w:hAnsiTheme'),
                 qn('w:eastAsiaTheme'), qn('w:cstheme')]:
        rFonts.attrib.pop(attr, None)


def setup_styles(doc):
    """设置文档样式：标题和正文字体规范

    【字体名称固定，禁止修改】
    '宋体' 是 OOXML w:eastAsia 标准映射名，Word 跨平台自动解析：
      Windows → SimSun  |  macOS → Songti SC
    禁止替换为 SimSun（仅 Windows）或 Songti SC（仅 macOS）
    '黑体' 同理：Windows → SimHei，macOS → Heiti SC
    """
    _apply_font(doc.styles['Title'],     22, bold=True,  cn_font='黑体')   # 封面：二号黑体加粗
    _apply_font(doc.styles['Heading 1'], 16, bold=True,  cn_font='宋体')   # 一级：三号宋体加粗
    _apply_font(doc.styles['Heading 2'], 15, bold=True,  cn_font='宋体')   # 二级：小三号宋体加粗
    _apply_font(doc.styles['Heading 3'], 14, bold=True,  cn_font='宋体')   # 三级：四号宋体加粗
    _apply_font(doc.styles['Heading 4'], 13, bold=True,  cn_font='宋体')   # 四级：小四号宋体加粗
    _apply_font(doc.styles['Heading 5'], 12, bold=False, cn_font='宋体')   # 五级：小四号宋体
    _apply_font(doc.styles['Normal'],    12, bold=False, cn_font='宋体')   # 正文：小四宋体

    # 2026 标书规范：行距 / 段距 / 首行缩进
    normal_pf = doc.styles['Normal'].paragraph_format
    normal_pf.line_spacing = 1.5
    normal_pf.space_before = Pt(6)         # 段前 0.5 行 ≈ 6pt（基于 12pt 行高）
    normal_pf.space_after = Pt(6)
    normal_pf.first_line_indent = Pt(24)   # 首行缩进 2 字符（小四 12pt × 2）
    for h in range(1, 6):
        h_pf = doc.styles[f'Heading {h}'].paragraph_format
        h_pf.line_spacing = 2.0            # 标题行距 2 倍
        h_pf.space_before = Pt(6)
        h_pf.space_after = Pt(6)
        h_pf.first_line_indent = Pt(0)     # 标题不缩进


def setup_multilevel_list(doc):
    """为 Heading 1-5 绑定 Word 多级列表自动编号（1, 1.1, 1.1.1, ...）。

    必须在 setup_styles 之后调用。每次调用增量追加 abstractNum + num，
    所以重复调用安全（但通常仅 create_document 时调一次）。

    架构对齐 solution-master/skills/go/scripts/docx_writer.py::setup_heading_numbering。
    """
    try:
        numbering_part = doc.part.numbering_part
    except Exception:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.parts.numbering import NumberingPart
        numbering_part = NumberingPart.new()
        doc.part.relate_to(numbering_part, RT.NUMBERING)

    numbering_elem = numbering_part.element

    existing_ids = [
        int(an.get(qn('w:abstractNumId')))
        for an in numbering_elem.findall(qn('w:abstractNum'))
    ]
    abstract_num_id = str(max(existing_ids) + 1) if existing_ids else '0'

    abstract_num = OxmlElement('w:abstractNum')
    abstract_num.set(qn('w:abstractNumId'), abstract_num_id)
    multi_level_type = OxmlElement('w:multiLevelType')
    multi_level_type.set(qn('w:val'), 'multilevel')
    abstract_num.append(multi_level_type)

    # 5 级编号格式：H1=1. / H2=1.1 / H3=1.1.1 / H4=1.1.1.1 / H5=1.1.1.1.1
    level_formats = [
        ('%1.', '0'),
        ('%1.%2', '1'),
        ('%1.%2.%3', '2'),
        ('%1.%2.%3.%4', '3'),
        ('%1.%2.%3.%4.%5', '4'),
    ]
    for lvl_text, ilvl in level_formats:
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), ilvl)
        start = OxmlElement('w:start');         start.set(qn('w:val'), '1');         lvl.append(start)
        num_fmt = OxmlElement('w:numFmt');      num_fmt.set(qn('w:val'), 'decimal'); lvl.append(num_fmt)
        lvl_text_elem = OxmlElement('w:lvlText'); lvl_text_elem.set(qn('w:val'), lvl_text); lvl.append(lvl_text_elem)
        lvl_jc = OxmlElement('w:lvlJc');        lvl_jc.set(qn('w:val'), 'left');    lvl.append(lvl_jc)
        suff = OxmlElement('w:suff');           suff.set(qn('w:val'), 'space');     lvl.append(suff)
        abstract_num.append(lvl)

    numbering_elem.append(abstract_num)

    existing_num_ids = [
        int(n.get(qn('w:numId')))
        for n in numbering_elem.findall(qn('w:num'))
    ]
    num_id = str(max(existing_num_ids) + 1) if existing_num_ids else '1'

    num_elem = OxmlElement('w:num')
    num_elem.set(qn('w:numId'), num_id)
    abstract_num_id_ref = OxmlElement('w:abstractNumId')
    abstract_num_id_ref.set(qn('w:val'), abstract_num_id)
    num_elem.append(abstract_num_id_ref)
    numbering_elem.append(num_elem)

    # 把 Heading 1-5 样式 bind 到 numId
    for h_level in range(1, 6):
        try:
            style = doc.styles[f'Heading {h_level}']
        except KeyError:
            continue
        pPr = style.element.get_or_add_pPr()
        old_numPr = pPr.find(qn('w:numPr'))
        if old_numPr is not None:
            pPr.remove(old_numPr)
        numPr = OxmlElement('w:numPr')
        ilvl_elem = OxmlElement('w:ilvl');   ilvl_elem.set(qn('w:val'), str(h_level - 1)); numPr.append(ilvl_elem)
        numId_elem = OxmlElement('w:numId'); numId_elem.set(qn('w:val'), num_id);          numPr.append(numId_elem)
        pPr.append(numPr)


def add_toc_field(doc, levels=4):
    """在文档当前位置插入 Word TOC 域：`TOC \\o "1-N" \\h \\z \\u`。

    打开 Word 后按 Ctrl+A → F9 更新即可显示目录树。
    """
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' TOC \\o "1-{levels}" \\h \\z \\u '
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar_begin)
    run._element.append(instrText)
    run._element.append(fldChar_separate)
    run._element.append(fldChar_end)
    return p


def validate_heading_hierarchy(doc, expected_outline):
    """验证 docx 中所有 Heading 段是否符合 expected_outline 结构。

    Args:
        doc: python-docx Document
        expected_outline: list[dict]，每项 {'numbering': '1.3.1', 'title': '...', 'depth': 3}

    Returns:
        (ok: bool, errors: list[str])
    """
    errors = []
    headings_in_doc = []
    for p in doc.paragraphs:
        if not p.style or not p.style.name.startswith('Heading'):
            continue
        try:
            level = int(p.style.name.split()[-1])
        except (ValueError, IndexError):
            continue
        headings_in_doc.append({'level': level, 'text': p.text.strip()})

    expected_titles = [(item['depth'], item['title'].strip()) for item in expected_outline]

    # 检查每个 expected outline item 都在 doc 中按顺序出现，level 匹配
    doc_idx = 0
    for exp_depth, exp_title in expected_titles:
        # 跳过非匹配 heading（doc 可能有额外 H1 章标题等）
        found = False
        while doc_idx < len(headings_in_doc):
            h = headings_in_doc[doc_idx]
            if h['text'] == exp_title:
                if h['level'] != exp_depth:
                    errors.append(
                        f"标题 '{exp_title}' 期望 Heading {exp_depth}，实际 Heading {h['level']}"
                    )
                doc_idx += 1
                found = True
                break
            doc_idx += 1
        if not found:
            errors.append(f"大纲标题 '{exp_title}' (depth {exp_depth}) 在 docx 中未找到")

    return (len(errors) == 0, errors)


def strip_numbering_prefix(text):
    """剥除文本开头的编号前缀（如 '1.3.1 招标方需求理解' / '1. 技术部分' → 标题正文）。

    兜底防御：subagent 偶尔会把编号写进 heading text，运行时强制清理。
    支持 'N.' / 'N.M' / 'N.M.K' / ... 末尾可选 '.'。
    """
    return re.sub(r'^\d+(\.\d+)*\.?\s+', '', text.strip())


def clean_doc_defaults(doc):
    """清除 python-docx 默认模板的主题字体引用

    python-docx 默认模板的 w:docDefaults 含主题字体引用，优先级最高，必须清除。
    """
    styles_elem = doc.styles.element
    doc_defaults = styles_elem.find(qn('w:docDefaults'))
    if doc_defaults is None:
        return
    rPr_default = doc_defaults.find(qn('w:rPrDefault'))
    if rPr_default is None:
        return
    rPr = rPr_default.find(qn('w:rPr'))
    if rPr is None:
        return
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        return
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:cs'), '宋体')
    for attr in [qn('w:asciiTheme'), qn('w:hAnsiTheme'),
                 qn('w:eastAsiaTheme'), qn('w:cstheme')]:
        rFonts.attrib.pop(attr, None)


# ── 运行级字体设置 ──────────────────────────────────

def apply_run_font(run, cn_font='宋体', latin_font='Times New Roman'):
    """设置单个 run 的字体（每个 run 也可能继承主题字体，必须逐 run 清除）"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), latin_font)
    rFonts.set(qn('w:hAnsi'), latin_font)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:cs'), cn_font)
    for attr in [qn('w:asciiTheme'), qn('w:hAnsiTheme'),
                 qn('w:eastAsiaTheme'), qn('w:cstheme')]:
        rFonts.attrib.pop(attr, None)


# ── 包装函数（替代 doc.add_heading / doc.add_paragraph）───

def add_heading_cn(doc, text, level, cn_font='宋体'):
    """添加中文标题，自动设置字体。

    text 中若含 '1.3.1 ' 前缀会被剥除——多级列表会自动加编号，
    残留编号会与自动编号叠加成 '1.3.1 1.3.1 标题'。
    """
    text = strip_numbering_prefix(text)
    h = doc.add_heading(text, level)
    for run in h.runs:
        apply_run_font(run, cn_font=cn_font)
    return h


def add_para_cn(doc, text, cn_font='宋体'):
    """添加中文段落（向后兼容，推荐使用 write_markdown 替代）"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        apply_run_font(run, cn_font=cn_font)
    return p


# ── Markdown 渲染函数 ──────────────────────────────────

def _apply_inline(p, text, cn_font='宋体'):
    """解析行内 **bold** *italic* `code`，向段落 p 追加 runs"""
    token_re = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    parts = token_re.split(text)
    for part in parts:
        if not part:
            continue
        run = p.add_run(part)
        if part.startswith('**') and part.endswith('**'):
            run.bold = True
            run.text = part[2:-2]
        elif part.startswith('*') and part.endswith('*'):
            run.italic = True
            run.text = part[1:-1]
        elif part.startswith('`') and part.endswith('`'):
            run.font.name = 'Courier New'
            run.text = part[1:-1]
        apply_run_font(run, cn_font=cn_font)


def write_markdown(doc, md_text, cn_font='宋体'):
    """将 Markdown 字符串渲染为 python-docx 内容。

    支持：标题(#)、无序列表(-/*)、有序列表(1.)、引用(>)、
    表格(|col|col|)、行内粗体(**)、斜体(*)、代码(``)、普通段落。
    替代 add_para_cn()，作为所有正文写入的统一入口。
    """
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue
        # 表格块检测
        if stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            def is_sep(row):
                return bool(re.match(r'^[\|\s\-:]+$', row))
            rows = [r for r in table_lines if not is_sep(r)]
            if rows:
                parsed = [[c.strip() for c in r.strip('|').split('|')] for r in rows]
                max_cols = max(len(r) for r in parsed)
                tbl = doc.add_table(rows=len(parsed), cols=max_cols)
                tbl.style = 'Table Grid'
                for r_i, row in enumerate(parsed):
                    for c_i in range(max_cols):
                        txt = row[c_i] if c_i < len(row) else ''
                        cell = tbl.cell(r_i, c_i)
                        cell.text = ''
                        p = cell.paragraphs[0]
                        if r_i == 0:
                            run = p.add_run(txt)
                            run.bold = True
                            apply_run_font(run, cn_font=cn_font)
                        else:
                            _apply_inline(p, txt, cn_font=cn_font)
            continue
        # 标题（# 数量 → Heading level；level=0 会渲染为 Title 样式，clamp 到 1）
        m = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if m:
            level = max(1, min(5, len(m.group(1))))
            add_heading_cn(doc, m.group(2), level, cn_font=cn_font)
            i += 1
            continue
        # 图片引用 ![caption](path)
        img_m = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$', stripped)
        if img_m:
            caption = img_m.group(1).strip() or None
            img_path = img_m.group(2).strip()
            add_picture_cn(doc, img_path, caption=caption, cn_font=cn_font)
            i += 1
            continue
        # 无序列表
        m = re.match(r'^[-*]\s+(.+)$', stripped)
        if m:
            p = doc.add_paragraph(style='Normal')
            bullet = p.add_run('• ')
            apply_run_font(bullet, cn_font=cn_font)
            _apply_inline(p, m.group(1), cn_font=cn_font)
            i += 1
            continue
        # 有序列表
        m = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if m:
            p = doc.add_paragraph(style='Normal')
            num = p.add_run(m.group(1) + '. ')
            apply_run_font(num, cn_font=cn_font)
            _apply_inline(p, m.group(2), cn_font=cn_font)
            i += 1
            continue
        # 引用
        m = re.match(r'^>\s*(.+)$', stripped)
        if m:
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.left_indent = Cm(0.75)
            _apply_inline(p, m.group(1), cn_font=cn_font)
            i += 1
            continue
        # 普通段落
        p = doc.add_paragraph(style='Normal')
        _apply_inline(p, stripped, cn_font=cn_font)
        i += 1


# ── 图片插入函数 ────────────────────────────────────

def add_picture_cn(doc, img_path, caption=None, width_cm=14.0, cn_font='宋体'):
    """插入图片并附带图题（caption 在图下方居中，小五宋体）。

    caption 文字应已是 '图 X-Y 说明' 格式（由调用方拼好），本函数不再二次加工。
    img_path 须为绝对路径或相对于工作目录的路径；文件不存在时降级为占位符。
    """
    if not os.path.exists(img_path):
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run('[图片占位符 — ' + (caption or img_path) + '：请手动插入图片]')
        run.italic = True
        apply_run_font(run, cn_font=cn_font)
        return
    doc.add_picture(img_path, width=Cm(width_cm))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    last_para.paragraph_format.first_line_indent = Pt(0)
    if caption:
        cap_p = doc.add_paragraph(style='Normal')
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.first_line_indent = Pt(0)
        cap_run = cap_p.add_run(caption)
        cap_run.font.size = Pt(9)  # 小五宋体（≈ 9pt）
        apply_run_font(cap_run, cn_font=cn_font)


# ── 文档创建快捷函数 ──────────────────────────────────

def create_document():
    """创建预配置的 DOCX 文档（A4、标准字体、页边距）

    返回已设置好样式和页面的 Document 对象，可直接写入内容。
    """
    doc = Document()
    setup_styles(doc)
    clean_doc_defaults(doc)
    setup_multilevel_list(doc)

    # 页面设置：A4，上下 2.5cm，左右 2.4cm
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    return doc
