# 投标文件大纲 DOCX 生成模板
# 由 taa skill Phase 2 调用，根据 Part B 大纲内容动态适配
#
# 使用方法：
# 1. AI 读取此模板，理解字体函数和页面设置
# 2. 根据 Part B 的实际章节内容，动态生成完整脚本
# 3. 执行生成的脚本，输出 .docx 文件
#
# 【重要】所有字符串参数必须用单引号 '...'，禁止双引号 "..."
# 原因：中文内容中常含 " 或 " 字符，在双引号字符串中会触发 SyntaxError
# （Python 3.12+ 对 Unicode 定界符更严格，U+201C/U+201D 均视为定界符）

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, datetime as dt

# ── 样式级字体设置 ──────────────────────────────────
def setup_styles(doc):
    def apply_font(style, size_pt, bold=False, cn_font='宋体', latin_font='Times New Roman'):
        style.font.size = Pt(size_pt)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.name = latin_font       # 高层 API 设置 ascii/hAnsi
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'),    latin_font)
        rFonts.set(qn('w:hAnsi'),    latin_font)
        rFonts.set(qn('w:eastAsia'), cn_font)
        rFonts.set(qn('w:cs'),       cn_font)
        # 删除主题字体引用（优先级高于明确字体名，必须清除）
        for attr in [qn('w:asciiTheme'), qn('w:hAnsiTheme'),
                     qn('w:eastAsiaTheme'), qn('w:cstheme')]:
            rFonts.attrib.pop(attr, None)
    # 【字体名称固定，禁止修改】
    # '宋体' 是 OOXML w:eastAsia 标准映射名，Word 跨平台自动解析：
    #   Windows → SimSun  |  macOS → Songti SC
    # 禁止替换为 SimSun（仅 Windows）或 Songti SC（仅 macOS）
    # '黑体' 同理：Windows → SimHei，macOS → Heiti SC
    apply_font(doc.styles['Title'],     22, bold=True,  cn_font='黑体')   # 封面：二号黑体加粗
    apply_font(doc.styles['Heading 1'], 16, bold=True,  cn_font='宋体')   # 一级：三号宋体加粗
    apply_font(doc.styles['Heading 2'], 15, bold=True,  cn_font='宋体')   # 二级：小三号宋体加粗
    apply_font(doc.styles['Heading 3'], 14, bold=True,  cn_font='宋体')   # 三级：四号宋体加粗
    apply_font(doc.styles['Normal'],    12, bold=False, cn_font='宋体')   # 正文：小四宋体

# ── 文档默认值字体清除 ─────────────────────────────────
# python-docx 默认模板的 w:docDefaults 含主题字体引用，优先级最高，必须清除
def clean_doc_defaults(doc):
    styles_elem = doc.styles.element
    doc_defaults = styles_elem.find(qn('w:docDefaults'))
    if doc_defaults is None:
        return
    rPr_default = doc_defaults.find(qn('w:rPrDefault'))
    if rPr_default is None:
        return
    rPr = rPr_default.find(qn('w:rPr'))
    if rPr is None:
        return
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        return
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:cs'), '宋体')
    for attr in [qn('w:asciiTheme'), qn('w:hAnsiTheme'),
                 qn('w:eastAsiaTheme'), qn('w:cstheme')]:
        rFonts.attrib.pop(attr, None)

# ── 运行级字体设置 ──────────────────────────────────
# 每个 run 也可能继承主题字体，必须逐 run 清除并设置
def apply_run_font(run, cn_font='宋体', latin_font='Times New Roman'):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), latin_font)
    rFonts.set(qn('w:hAnsi'), latin_font)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:cs'), cn_font)
    for attr in [qn('w:asciiTheme'), qn('w:hAnsiTheme'),
                 qn('w:eastAsiaTheme'), qn('w:cstheme')]:
        rFonts.attrib.pop(attr, None)

# ── 包装函数（替代 doc.add_heading / doc.add_paragraph）───
def add_heading_cn(doc, text, level, cn_font='宋体'):
    h = doc.add_heading(text, level)
    for run in h.runs:
        apply_run_font(run, cn_font=cn_font)
    return h

def add_para_cn(doc, text, cn_font='宋体'):
    p = doc.add_paragraph(text)
    for run in p.runs:
        apply_run_font(run, cn_font=cn_font)
    return p

# ── 主流程 ──────────────────────────────────────────

# 输出目录
output_dir = './output'
os.makedirs(output_dir, exist_ok=True)

doc = Document()

# 页面设置（A4，上下2.5cm，左右2.4cm）
sec = doc.sections[0]
sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(2.4); sec.right_margin = Cm(2.4)

setup_styles(doc)
clean_doc_defaults(doc)  # 清除文档默认值中的主题字体

# 封面（黑体）
add_heading_cn(doc, '[项目名称]', 0, cn_font='黑体')
add_para_cn(doc, '投标文件大纲')
add_para_cn(doc, f'生成时间：{dt.date.today()}')
add_para_cn(doc, '{VENDOR_NAME}')
doc.add_page_break()

# 目录占位
add_heading_cn(doc, '目录', 1)
add_para_cn(doc, '[Word 自动生成目录，按 Ctrl+A → F9 更新]')
doc.add_page_break()

# [按 Part B 大纲内容逐章节生成，示意如下]
# 一级标题（章）：add_heading_cn(doc, '一、技术部分', 1)
# 二级标题（节）：add_heading_cn(doc, '1.1 技术偏离表', 2)
# 二级标题（节）：add_heading_cn(doc, '1.2 项目理解与需求分析', 2)
# 三级标题：      add_heading_cn(doc, '1.2.1 招标方需求理解', 3)
# 正文：          add_para_cn(doc, '[请在此处填写本节内容]')

# ── 动态生成逻辑 ──────────────────────────────────
# 遍历 Part B 的每一行，根据章节编号判断级别：
# - 一、→ level 1 (add_heading_cn(doc, text, 1))  # 仅"一、技术部分"
# - 1.X → level 2 (add_heading_cn(doc, text, 2))
# - 1.X.Y → level 3 (add_heading_cn(doc, text, 3))
# 在每个章节后添加占位段落：add_para_cn(doc, '[请在此处填写本节内容]')
#
# 示例：若 Part B 包含以下内容（仅技术部分）：
# 一、技术部分
# 1.1 技术偏离表
# 1.2 项目理解与需求分析
# 1.2.1 招标方需求理解
#
# 则生成代码：
# add_heading_cn(doc, '一、技术部分', 1)
# add_para_cn(doc, '[请在此处填写本节内容]')
# add_heading_cn(doc, '1.1 技术偏离表', 2)
# add_para_cn(doc, '[请在此处填写本节内容]')
# add_heading_cn(doc, '1.2 项目理解与需求分析', 2)
# add_para_cn(doc, '[请在此处填写本节内容]')
# add_heading_cn(doc, '1.2.1 招标方需求理解', 3)
# add_para_cn(doc, '[请在此处填写本节内容]')

timestamp = dt.datetime.now().strftime('%Y%m%d_%H%M%S')
filename = os.path.join(output_dir, f'投标文件大纲_{timestamp}.docx')
doc.save(filename)
abs_path = os.path.abspath(filename)
print(f'大纲已保存：{abs_path}')
