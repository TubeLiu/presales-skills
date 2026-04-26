#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""TRV 审核后 DOCX 智能修订工具（v2.0 — 指令驱动）。

接收 Claude 生成的结构化修订指令 JSON，对 DOCX 执行通用修订操作。
工具本身不包含任何项目专属规则，所有修订逻辑由 AI 在审核阶段动态生成。

支持的指令类型：
  - paragraph_text_replace  段落内查找替换（保留 run 格式）
  - paragraph_full_replace  整段文本替换（格式降级到首 run）
  - table_cell_replace      按表头+行内容定位单元格替换
  - global_text_replace     全文查找替换（段落 + 表格）
  - paragraph_insert_after  在匹配段落后插入新段落
  - paragraph_delete        删除匹配段落

用法：
  python3 tools/trv_docx_reviser.py \\
    --input 技术标书.docx \\
    --instructions revision_instructions.json \\
    [--output-dir output/trv]
"""

from __future__ import annotations

import argparse
import copy
import json
import sys
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Normalization
# ---------------------------------------------------------------------------

def _normalize(text: str) -> str:
    """归一化文本：去除首尾空白、压缩连续空白为单空格。"""
    return ' '.join(text.strip().split())


# ---------------------------------------------------------------------------
# Result tracking
# ---------------------------------------------------------------------------

@dataclass
class InstructionResult:
    id: str
    status: str  # applied / not_found / error / skipped
    details: str = ''


@dataclass
class ReviseStats:
    total: int = 0
    applied: int = 0
    not_found: int = 0
    errors: int = 0
    skipped: int = 0
    results: list[InstructionResult] = field(default_factory=list)

    def record(self, result: InstructionResult) -> None:
        self.results.append(result)
        self.total += 1
        if result.status == 'applied':
            self.applied += 1
        elif result.status == 'not_found':
            self.not_found += 1
        elif result.status == 'error':
            self.errors += 1
        elif result.status == 'skipped':
            self.skipped += 1

    def to_dict(self) -> dict:
        return {
            'summary': {
                'total': self.total,
                'applied': self.applied,
                'not_found': self.not_found,
                'errors': self.errors,
                'skipped': self.skipped,
            },
            'results': [
                {'id': r.id, 'status': r.status, 'details': r.details}
                for r in self.results
            ],
        }


# ---------------------------------------------------------------------------
# Run-level text replacement (format-preserving)
# ---------------------------------------------------------------------------

def _replace_in_runs(paragraph, find_text: str, replace_text: str) -> bool:
    """在段落的 run 级别执行文本替换，尽量保留原有格式。

    策略：
    1. 拼接所有 run.text 得到完整段落文本
    2. 在拼接文本中定位 find_text
    3. 将替换文本分配到对应的 run 中，保留首个命中 run 的格式

    归一化匹配失败时返回 False（不降级为 para.text= 以避免破坏格式）。
    """
    if not find_text:
        return False

    runs = paragraph.runs
    if not runs:
        return False

    # 构建 run 文本与字符位置映射
    full_text = ''.join(r.text for r in runs)
    start = full_text.find(find_text)
    if start == -1:
        # 尝试归一化匹配：仅当归一化后文本长度不变时可安全映射回原文位置
        norm_full = _normalize(full_text)
        norm_find = _normalize(find_text)
        if len(norm_full) == len(full_text):
            # 空白未被压缩，偏移量可直接映射
            start = norm_full.find(norm_find)
            if start != -1:
                # 用原文偏移做 run 级替换（继续走下面的逻辑）
                find_text = full_text[start:start + len(norm_find)]
            else:
                return False
        else:
            # 归一化改变了长度，无法安全映射回原文位置，放弃本次匹配
            return False

    end = start + len(find_text)

    # 映射每个字符属于哪个 run
    char_to_run = []
    for i, run in enumerate(runs):
        char_to_run.extend([i] * len(run.text))

    if not char_to_run:
        return False

    first_run_idx = char_to_run[start]
    last_run_idx = char_to_run[end - 1] if end > 0 else first_run_idx

    # 计算每个 run 内的偏移
    run_starts = []
    pos = 0
    for run in runs:
        run_starts.append(pos)
        pos += len(run.text)

    # 单 run 内替换：最简单的情况
    if first_run_idx == last_run_idx:
        run = runs[first_run_idx]
        local_start = start - run_starts[first_run_idx]
        local_end = end - run_starts[first_run_idx]
        run.text = run.text[:local_start] + replace_text + run.text[local_end:]
        return True

    # 跨 run 替换：把替换文本放在第一个 run，清空中间 run，裁剪最后 run
    first_run = runs[first_run_idx]
    last_run = runs[last_run_idx]

    local_start_in_first = start - run_starts[first_run_idx]
    local_end_in_last = end - run_starts[last_run_idx]

    first_run.text = first_run.text[:local_start_in_first] + replace_text
    last_run.text = last_run.text[local_end_in_last:]

    # 清空中间 run
    for i in range(first_run_idx + 1, last_run_idx):
        runs[i].text = ''

    return True


def _replace_all_in_runs(paragraph, find_text: str, replace_text: str) -> int:
    """在段落中替换所有出现的 find_text，返回替换次数。"""
    if not find_text or find_text == replace_text:
        return 0

    # 若 replace_text 包含 find_text，迭代替换会无限循环
    # 此时退化为单次 paragraph.text 替换
    if find_text in replace_text:
        full_text = ''.join(r.text for r in paragraph.runs) if paragraph.runs else paragraph.text
        if find_text not in full_text:
            return 0
        count = full_text.count(find_text)
        new_text = full_text.replace(find_text, replace_text)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ''
        else:
            paragraph.text = new_text
        return count

    count = 0
    max_iterations = 100
    for _ in range(max_iterations):
        if _replace_in_runs(paragraph, find_text, replace_text):
            count += 1
        else:
            break
    return count


# ---------------------------------------------------------------------------
# Paragraph matching
# ---------------------------------------------------------------------------

def _match_paragraph(paragraph, match_spec: dict) -> bool:
    """判断段落是否匹配给定的 match 规范。"""
    method = match_spec.get('method', 'contains')
    text = paragraph.text

    if method == 'contains':
        target = match_spec.get('text', '')
        if target not in text and _normalize(target) not in _normalize(text):
            return False
        # 可选的同段落上下文校验（用于消歧，检查同一段落内是否包含该文本）
        ctx_before = match_spec.get('context_before')
        if ctx_before and ctx_before not in text and _normalize(ctx_before) not in _normalize(text):
            return False
        return True

    elif method == 'exact':
        target = match_spec.get('text', '')
        return _normalize(text) == _normalize(target)

    elif method == 'regex':
        import re
        pattern = match_spec.get('pattern', '')
        return bool(re.search(pattern, text))

    return False


# ---------------------------------------------------------------------------
# Table cell matching
# ---------------------------------------------------------------------------

def _find_table_cell(doc: Document, match_spec: dict):
    """按表头内容 + 行标识定位表格单元格。

    match_spec 结构：
    {
        "method": "table_header_row",
        "header_contains": ["序号", "投标规格"],
        "row_match": {"column": "序号", "value": "15"},
        "target_column": "投标规格"
    }
    """
    header_contains = match_spec.get('header_contains', [])
    row_match = match_spec.get('row_match', {})
    target_column = match_spec.get('target_column', '')

    for table in doc.tables:
        if not table.rows:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]

        # 检查表头是否包含所有要求的列名
        if not all(any(h in header for header in headers) for h in header_contains):
            continue

        # 查找目标列和行标识列的索引
        target_col_idx = None
        row_id_col_idx = None
        row_id_column = row_match.get('column', '')
        for i, header in enumerate(headers):
            if target_column and target_column in header:
                target_col_idx = i
            if row_id_column and row_id_column in header:
                row_id_col_idx = i

        if target_col_idx is None:
            continue

        # 如果没有 row_match，返回第一个数据行的目标列
        if not row_match or row_id_col_idx is None:
            if len(table.rows) > 1:
                return table.rows[1].cells[target_col_idx]
            continue

        # 按行标识值查找
        row_id_value = str(row_match.get('value', ''))
        for row in table.rows[1:]:
            if len(row.cells) > max(target_col_idx, row_id_col_idx):
                if row.cells[row_id_col_idx].text.strip() == row_id_value:
                    return row.cells[target_col_idx]

    return None


# ---------------------------------------------------------------------------
# Instruction executors
# ---------------------------------------------------------------------------

def _apply_paragraph_text_replace(doc: Document, instruction: dict, stats: ReviseStats) -> None:
    """段落内查找替换（保留 run 格式）。"""
    inst_id = instruction['id']
    match_spec = instruction.get('match', {})
    action = instruction.get('action', {})
    find_text = action.get('find', '')
    replace_text = action.get('replace', '')

    if not find_text:
        stats.record(InstructionResult(inst_id, 'error', 'action.find is empty'))
        return

    applied = False
    for para in doc.paragraphs:
        if _match_paragraph(para, match_spec):
            if _replace_in_runs(para, find_text, replace_text):
                applied = True
                break  # 默认只替换第一个匹配

    if applied:
        stats.record(InstructionResult(inst_id, 'applied', f'replaced in paragraph'))
    else:
        stats.record(InstructionResult(inst_id, 'not_found', f'match text not found: {find_text[:50]}'))


def _apply_paragraph_full_replace(doc: Document, instruction: dict, stats: ReviseStats) -> None:
    """整段文本替换（格式降级到首 run）。"""
    inst_id = instruction['id']
    match_spec = instruction.get('match', {})
    action = instruction.get('action', {})
    new_text = action.get('new_text', '')

    applied = False
    for para in doc.paragraphs:
        if _match_paragraph(para, match_spec):
            # 保留首 run 格式，清空其余
            if para.runs:
                para.runs[0].text = new_text
                for run in para.runs[1:]:
                    run.text = ''
            else:
                para.text = new_text
            applied = True
            break

    if applied:
        stats.record(InstructionResult(inst_id, 'applied', 'paragraph fully replaced'))
    else:
        stats.record(InstructionResult(inst_id, 'not_found', 'matching paragraph not found'))


def _apply_table_cell_replace(doc: Document, instruction: dict, stats: ReviseStats) -> None:
    """按表头+行内容定位单元格替换。"""
    inst_id = instruction['id']
    match_spec = instruction.get('match', {})
    action = instruction.get('action', {})
    find_text = action.get('find', '')
    replace_text = action.get('replace', '')

    cell = _find_table_cell(doc, match_spec)
    if cell is None:
        stats.record(InstructionResult(inst_id, 'not_found', 'table cell not found'))
        return

    if find_text:
        # 在单元格内查找替换
        original = cell.text
        if find_text in original:
            # 尝试在段落 run 级别替换
            replaced = False
            for para in cell.paragraphs:
                if find_text in para.text:
                    if _replace_in_runs(para, find_text, replace_text):
                        replaced = True
                        break
            if not replaced:
                cell.text = original.replace(find_text, replace_text)
            stats.record(InstructionResult(inst_id, 'applied', 'table cell text replaced'))
        elif _normalize(find_text) in _normalize(original):
            # 归一化匹配成功但精确匹配失败，直接用 cell.text 替换（会丢失格式，但表格单元格格式通常简单）
            normalized = _normalize(original)
            cell.text = normalized.replace(_normalize(find_text), replace_text)
            stats.record(InstructionResult(inst_id, 'applied', 'table cell text replaced (normalized, formatting may change)'))
        else:
            stats.record(InstructionResult(inst_id, 'not_found', f'find text not in cell: {find_text[:50]}'))
    else:
        # 无 find_text 时直接设置整个单元格
        new_text = action.get('replace', action.get('new_text', ''))
        cell.text = new_text
        stats.record(InstructionResult(inst_id, 'applied', 'table cell fully replaced'))


def _apply_global_text_replace(doc: Document, instruction: dict, stats: ReviseStats) -> None:
    """全文查找替换（段落 + 表格单元格）。"""
    inst_id = instruction['id']
    action = instruction.get('action', {})
    find_text = action.get('find', '')
    replace_text = action.get('replace', '')

    if not find_text:
        stats.record(InstructionResult(inst_id, 'error', 'action.find is empty'))
        return

    total_count = 0

    # 段落替换
    for para in doc.paragraphs:
        count = _replace_all_in_runs(para, find_text, replace_text)
        total_count += count

    # 表格单元格替换
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    count = _replace_all_in_runs(para, find_text, replace_text)
                    total_count += count

    if total_count > 0:
        stats.record(InstructionResult(inst_id, 'applied', f'replaced {total_count} occurrences'))
    else:
        stats.record(InstructionResult(inst_id, 'not_found', f'text not found anywhere: {find_text[:50]}'))


def _apply_paragraph_insert_after(doc: Document, instruction: dict, stats: ReviseStats) -> None:
    """在匹配段落后插入新段落。"""
    inst_id = instruction['id']
    match_spec = instruction.get('match', {})
    action = instruction.get('action', {})
    insert_text = action.get('insert_text', '')

    if not insert_text:
        stats.record(InstructionResult(inst_id, 'error', 'action.insert_text is empty'))
        return

    for para in doc.paragraphs:
        if _match_paragraph(para, match_spec):
            # 在段落后插入新段落
            new_para = copy.deepcopy(para._element)
            # 清空内容，设置新文本
            for child in list(new_para):
                if child.tag.endswith('}r'):
                    new_para.remove(child)
            run_elem = copy.deepcopy(para.runs[0]._element) if para.runs else None
            if run_elem is not None:
                # 清空 run 文本，设置新文本
                for t_elem in run_elem.findall(qn('w:t')):
                    run_elem.remove(t_elem)
                from docx.oxml import OxmlElement
                t = OxmlElement('w:t')
                t.text = insert_text
                t.set(qn('xml:space'), 'preserve')
                run_elem.append(t)
                new_para.append(run_elem)
            else:
                from docx.oxml import OxmlElement
                run_elem = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = insert_text
                t.set(qn('xml:space'), 'preserve')
                run_elem.append(t)
                new_para.append(run_elem)
            para._element.addnext(new_para)
            stats.record(InstructionResult(inst_id, 'applied', 'paragraph inserted'))
            return

    stats.record(InstructionResult(inst_id, 'not_found', 'anchor paragraph not found'))


def _apply_paragraph_delete(doc: Document, instruction: dict, stats: ReviseStats) -> None:
    """删除匹配的段落。"""
    inst_id = instruction['id']
    match_spec = instruction.get('match', {})

    for para in doc.paragraphs:
        if _match_paragraph(para, match_spec):
            parent = para._element.getparent()
            if parent is not None:
                parent.remove(para._element)
                stats.record(InstructionResult(inst_id, 'applied', 'paragraph deleted'))
                return

    stats.record(InstructionResult(inst_id, 'not_found', 'matching paragraph not found for deletion'))


# ---------------------------------------------------------------------------
# Instruction dispatcher
# ---------------------------------------------------------------------------

_INSTRUCTION_HANDLERS = {
    'paragraph_text_replace': _apply_paragraph_text_replace,
    'paragraph_full_replace': _apply_paragraph_full_replace,
    'table_cell_replace': _apply_table_cell_replace,
    'global_text_replace': _apply_global_text_replace,
    'paragraph_insert_after': _apply_paragraph_insert_after,
    'paragraph_delete': _apply_paragraph_delete,
}


def apply_instructions(doc: Document, instructions: list, stats: ReviseStats) -> None:
    """按顺序执行所有修订指令。"""
    for instruction in instructions:
        inst_id = instruction.get('id', 'unknown')
        inst_type = instruction.get('type', '')
        handler = _INSTRUCTION_HANDLERS.get(inst_type)
        if handler is None:
            stats.record(InstructionResult(inst_id, 'error', f'unknown instruction type: {inst_type}'))
            continue
        try:
            handler(doc, instruction, stats)
        except Exception as e:
            stats.record(InstructionResult(inst_id, 'error', str(e)))


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def revise_docx(input_path: Path, instructions_path: Path, output_dir: Path) -> dict:
    """执行指令驱动的 DOCX 修订。"""
    # 加载指令
    instructions_data = json.loads(instructions_path.read_text(encoding='utf-8'))
    instructions = instructions_data.get('instructions', [])

    output_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    output_path = output_dir / f'{input_path.stem}_修订版_{timestamp}.docx'

    doc = Document(str(input_path))
    stats = ReviseStats()

    apply_instructions(doc, instructions, stats)

    doc.save(str(output_path))

    return {
        'output_file': str(output_path),
        **stats.to_dict(),
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description='TRV 审核后 DOCX 智能修订工具（指令驱动）')
    parser.add_argument('--input', required=True, help='待修订 DOCX 文件路径')
    parser.add_argument('--instructions', required=True, help='修订指令 JSON 文件路径')
    parser.add_argument('--output-dir', default='output/trv', help='修订版输出目录')
    # 废弃参数：接受但忽略，保持向后兼容
    parser.add_argument('--report', help='[废弃] 已被 --instructions 替代')
    parser.add_argument('--review-type', help='[废弃] 不再需要')
    parser.add_argument('--scope', help='[废弃] 修订范围由指令 JSON 控制')
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    input_path = Path(args.input)
    instructions_path = Path(args.instructions)
    output_dir = Path(args.output_dir)

    # 废弃参数提示
    if args.report or args.review_type or args.scope:
        print('[warn] --report, --review-type, --scope 已废弃，请使用 --instructions 传入修订指令 JSON',
              file=sys.stderr)

    if input_path.suffix.lower() != '.docx':
        print(json.dumps({'status': 'skipped', 'reason': 'input is not docx'}, ensure_ascii=False))
        return 0
    if not input_path.exists():
        raise FileNotFoundError(f'input file not found: {input_path}')
    if not instructions_path.exists():
        raise FileNotFoundError(f'instructions file not found: {instructions_path}')

    result = revise_docx(input_path, instructions_path, output_dir)
    print(json.dumps({'status': 'ok', **result}, ensure_ascii=False))
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
