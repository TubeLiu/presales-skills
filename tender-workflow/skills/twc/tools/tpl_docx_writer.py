#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
tpl DOCX 渲染工具

从 JSON 数据文件读取结构化文档内容，渲染为格式规范的 DOCX 文件。
内容与渲染逻辑分离，避免在 Python 脚本中内嵌大量中文字符串。

用法：
  python3 tools/tpl_docx_writer.py <content.json> <output.docx>

JSON 数据格式：
  {
    "cover": {
      "title": "项目名称",
      "subtitle": "技术规格与评标办法",
      "info": ["行业类型：通用企业", "项目预算：100 万元", "编制日期：2026 年 04 月"]
    },
    "parts": [
      {
        "title": "第一部分  技术规格与要求",
        "sections": [
          {
            "heading": "1  项目总体说明",
            "level": 2,
            "children": [
              {"heading": "1.1  项目背景", "level": 3, "body": ["段落文本..."]},
              {"heading": "1.2  建设目标", "level": 3, "bullets": ["条目1", "条目2"]}
            ]
          },
          {
            "heading": "3  性能需求规格",
            "level": 2,
            "body": ["段落文本..."],
            "tables": [
              {
                "headers": ["列1", "列2", "列3"],
                "rows": [["a", "b", "c"], ["d", "e", "f"]]
              }
            ]
          }
        ]
      }
    ]
  }
"""

import json
import sys
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 字体设置 ──────────────────────────────────────────

def _set_rfonts(rPr, cn, en):
    rF = rPr.find(qn('w:rFonts'))
    if rF is None:
        rF = OxmlElement('w:rFonts')
        rPr.insert(0, rF)
    rF.set(qn('w:ascii'), en)
    rF.set(qn('w:hAnsi'), en)
    rF.set(qn('w:eastAsia'), cn)
    rF.set(qn('w:cs'), en)
    for attr in [qn('w:asciiTheme'), qn('w:hAnsiTheme'),
                 qn('w:eastAsiaTheme'), qn('w:cstheme')]:
        rF.attrib.pop(attr, None)


def apply_run_font(run, cn='仿宋_GB2312', en='Times New Roman', size=Pt(12)):
    run.font.size = size
    run.font.name = en
    _set_rfonts(run._element.get_or_add_rPr(), cn, en)


def setup_styles(doc):
    # 清除 docDefaults 中的主题字体
    styles_el = doc.styles.element
    doc_defaults = styles_el.find(qn('w:docDefaults'))
    if doc_defaults is not None:
        rPr_default = doc_defaults.find(qn('w:rPrDefault'))
        if rPr_default is not None:
            rPr = rPr_default.find(qn('w:rPr'))
            if rPr is not None:
                rF = rPr.find(qn('w:rFonts'))
                if rF is not None:
                    _set_rfonts(rPr, '仿宋_GB2312', 'Times New Roman')

    s = doc.styles['Normal']
    s.font.size = Pt(12)
    s.font.name = 'Times New Roman'
    pf = s.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    _set_rfonts(s.element.get_or_add_rPr(), '仿宋_GB2312', 'Times New Roman')

    for i in range(1, 5):
        hs = doc.styles['Heading %d' % i]
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.font.bold = True


# ── 内容写入函数 ──────────────────────────────────────

def add_heading(doc, text, level=1, cn='黑体', size=None):
    sizes = {1: Pt(22), 2: Pt(16), 3: Pt(14), 4: Pt(12)}
    if size is None:
        size = sizes.get(level, Pt(12))
    h = doc.add_heading(level=level)
    r = h.add_run(text)
    apply_run_font(r, cn, 'Times New Roman', size)
    r.bold = True
    return h


def add_para(doc, text, bold=False, indent=True, size=Pt(12)):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    r = p.add_run(text)
    apply_run_font(r, '仿宋_GB2312', 'Times New Roman', size)
    r.bold = bold
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    apply_run_font(r, '仿宋_GB2312', 'Times New Roman', Pt(12))
    return p


def add_table(doc, headers, rows):
    cols = len(headers)
    t = doc.add_table(rows=1, cols=cols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, txt in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        apply_run_font(r, '黑体', 'Times New Roman', Pt(10.5))
        r.bold = True
    for row_data in rows:
        row = t.add_row()
        bold = row_data.get('bold', False) if isinstance(row_data, dict) else False
        cells = row_data.get('cells', row_data) if isinstance(row_data, dict) else row_data
        for i, txt in enumerate(cells):
            if i >= cols:
                break
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(txt))
            apply_run_font(r, '仿宋_GB2312', 'Times New Roman', Pt(10.5))
            r.bold = bold
    doc.add_paragraph()
    return t


# ── 节渲染 ────────────────────────────────────────────

def render_section(doc, section):
    """递归渲染一个 section 节点"""
    if 'heading' in section:
        level = section.get('level', 2)
        font = section.get('font', '黑体')
        size = section.get('size', None)
        if size:
            size = Pt(size)
        add_heading(doc, section['heading'], level, font, size)

    for text in section.get('body', []):
        bold = False
        indent = True
        if isinstance(text, dict):
            bold = text.get('bold', False)
            indent = text.get('indent', True)
            text = text.get('text', '')
        add_para(doc, text, bold=bold, indent=indent)

    for text in section.get('bullets', []):
        add_bullet(doc, text)

    for tbl in section.get('tables', []):
        add_table(doc, tbl['headers'], tbl['rows'])

    for numbered in section.get('numbered', []):
        add_para(doc, numbered, indent=False)

    for child in section.get('children', []):
        render_section(doc, child)


# ── 主流程 ────────────────────────────────────────────

def render_docx(data, output_path):
    doc = Document()
    setup_styles(doc)
    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.17)
        sec.right_margin = Cm(3.17)

    # 封面
    cover = data.get('cover', {})
    if cover:
        for _ in range(6):
            doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(cover.get('title', ''))
        apply_run_font(r, '方正小标宋简体', 'Times New Roman', Pt(26))
        r.bold = True
        doc.add_paragraph()
        if cover.get('subtitle'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(cover['subtitle'])
            apply_run_font(r, '黑体', 'Times New Roman', Pt(22))
            r.bold = True
        for _ in range(4):
            doc.add_paragraph()
        for line in cover.get('info', []):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line)
            apply_run_font(r, '仿宋_GB2312', 'Times New Roman', Pt(14))
        doc.add_page_break()

    # 目录占位
    add_heading(doc, '目  录', 1, '黑体', Pt(22))
    add_para(doc, '（请在 Word 中按 Ctrl+A，然后按 F9 更新目录）',
             indent=False, size=Pt(10.5))
    doc.add_page_break()

    # 各部分
    for part in data.get('parts', []):
        add_heading(doc, part['title'], 1, '方正小标宋简体', Pt(22))
        for section in part.get('sections', []):
            render_section(doc, section)
        doc.add_page_break()

    doc.save(output_path)
    return output_path


def main():
    if len(sys.argv) < 3:
        print('用法: python3 tools/tpl_docx_writer.py <content.json> <output.docx>')
        sys.exit(1)

    json_path = sys.argv[1]
    output_path = sys.argv[2]

    if not os.path.exists(json_path):
        print('ERROR: JSON 文件不存在: %s' % json_path, file=sys.stderr)
        sys.exit(1)

    with open(json_path, 'r', encoding='utf-8') as f:
        raw = f.read()
    # 清理 Write 工具传输时可能产生的 U+FFFD 字符碎片
    removed = raw.count('\ufffd')
    if removed:
        raw = raw.replace('\ufffd', '')
        print('JSON_CLEANED:removed=%d U+FFFD chars from %s' % (removed, json_path))
    data = json.loads(raw)

    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    render_docx(data, output_path)
    print('SUCCESS:%s' % output_path)


if __name__ == '__main__':
    main()
