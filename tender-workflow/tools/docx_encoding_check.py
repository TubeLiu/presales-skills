#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX 编码完整性检查与自动修复工具

检测并修复 DOCX 文件中因 UTF-8 字符截断导致的乱码（U+FFFD 替换字符）。

修复策略：直接删除 U+FFFD 字符。
UTF-8 截断导致的 U+FFFD 是字节碎片，不携带可恢复信息，删除后上下文
自然衔接。若删除后语义不通顺，属于生成源头问题，需重新生成文档。

用法：
  python3 tools/docx_encoding_check.py --check <file.docx>
  python3 tools/docx_encoding_check.py --fix <file.docx> [--output <out.docx>] [--max-retries 3]

退出码：
  0 = 无乱码（或已全部清除）
  1 = 存在乱码且未修复（--check 模式）
  2 = 重试后仍有残留乱码（--fix 模式）
"""

import sys
import os
import argparse


REPLACEMENT_CHAR = '\ufffd'


def fix_text_file(filepath, output_path=None):
    """修复纯文本/JSON 文件中的 U+FFFD 字符。

    返回 (removed_count, output_path)
    """
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    count = content.count(REPLACEMENT_CHAR)
    if count == 0:
        return 0, filepath
    content = content.replace(REPLACEMENT_CHAR, '')
    out = output_path or filepath
    with open(out, 'w', encoding='utf-8') as f:
        f.write(content)
    return count, out


def scan_docx(filepath):
    """扫描 DOCX 中所有含 U+FFFD 的文本位置。

    返回 (issues, total_paragraphs)
    issues: [{'location': str, 'text': str}, ...]
    """
    from docx import Document
    doc = Document(filepath)
    issues = []

    for i, para in enumerate(doc.paragraphs):
        if REPLACEMENT_CHAR in para.text:
            issues.append({
                'location': 'para %d' % (i + 1),
                'text': para.text.strip()[:80],
            })

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                if REPLACEMENT_CHAR in cell.text:
                    issues.append({
                        'location': 'table%d[%d][%d]' % (ti, ri, ci),
                        'text': cell.text.strip()[:80],
                    })

    return issues, len(doc.paragraphs)


def clean_runs(filepath, output_path=None):
    """清除 DOCX 中所有 run 内的 U+FFFD 字符，保存文件。

    返回清除的 U+FFFD 字符总数。
    """
    from docx import Document
    doc = Document(filepath)
    removed = 0

    for para in doc.paragraphs:
        for run in para.runs:
            count = run.text.count(REPLACEMENT_CHAR)
            if count:
                run.text = run.text.replace(REPLACEMENT_CHAR, '')
                removed += count

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        count = run.text.count(REPLACEMENT_CHAR)
                        if count:
                            run.text = run.text.replace(REPLACEMENT_CHAR, '')
                            removed += count

    doc.save(output_path or filepath)
    return removed


def cmd_check(args):
    issues, total = scan_docx(args.file)
    if not issues:
        print('ENCODING_OK:paragraphs=%d' % total)
        return 0

    print('ENCODING_ERROR:%d' % len(issues))
    for item in issues[:20]:
        print('  %s: %s' % (item['location'], item['text']))
    if len(issues) > 20:
        print('  ... and %d more' % (len(issues) - 20))
    return 1


def cmd_fix(args):
    max_retries = args.max_retries
    out = args.output or args.file

    # 首次检查
    issues, _ = scan_docx(args.file)
    if not issues:
        print('ENCODING_OK:no_issues')
        return 0

    total_removed = 0
    src = args.file

    for attempt in range(1, max_retries + 1):
        removed = clean_runs(src, out)
        total_removed += removed

        remaining, _ = scan_docx(out)
        if not remaining:
            print('FIX_OK:removed=%d,retries=%d,file=%s'
                  % (total_removed, attempt, out))
            return 0

        print('FIX_RETRY:%d/%d removed=%d,remaining=%d'
              % (attempt, max_retries, removed, len(remaining)))
        # 下一轮从修复后的文件继续
        src = out

    # 重试耗尽仍有残留
    remaining, _ = scan_docx(out)
    print('FIX_FAILED:removed=%d,remaining=%d,file=%s'
          % (total_removed, len(remaining), out))
    for item in remaining[:10]:
        print('  UNFIXED %s: %s' % (item['location'], item['text']))
    return 2


def main():
    parser = argparse.ArgumentParser(
        description='DOCX 编码完整性检查与自动修复')
    parser.add_argument('file', help='DOCX 文件路径')
    sub = parser.add_subparsers(dest='command')

    sub.add_parser('--check', help='仅检测')
    p_fix = sub.add_parser('--fix', help='检测并修复')
    p_fix.add_argument('--output', '-o', default=None,
                       help='输出路径（默认原地覆写）')
    p_fix.add_argument('--max-retries', type=int, default=3,
                       help='最大重试次数（默认 3）')

    # 兼容 --check / --fix 作为 flag 的旧用法
    parser.add_argument('--check', action='store_true', dest='flag_check',
                        help=argparse.SUPPRESS)
    parser.add_argument('--fix', action='store_true', dest='flag_fix',
                        help=argparse.SUPPRESS)
    parser.add_argument('--output', '-o', default=None, dest='flag_output',
                        help=argparse.SUPPRESS)
    parser.add_argument('--max-retries', type=int, default=3,
                        dest='flag_max_retries', help=argparse.SUPPRESS)

    args = parser.parse_args()

    if not os.path.exists(args.file):
        print('ERROR: file not found: %s' % args.file, file=sys.stderr)
        sys.exit(1)

    if args.flag_check:
        sys.exit(cmd_check(args))
    elif args.flag_fix:
        args.output = args.flag_output
        args.max_retries = args.flag_max_retries
        sys.exit(cmd_fix(args))
    else:
        parser.print_help()
        sys.exit(1)


if __name__ == '__main__':
    main()
