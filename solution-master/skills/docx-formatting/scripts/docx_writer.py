#!/usr/bin/env python3
"""
Solution Master DOCX 渲染器

将 Markdown 内容渲染为符合规范的 DOCX 文件。
支持标题（H1-H5）、正文、列表、表格、图片、引用。

字体规范：
  封面标题：22pt Bold 黑体 / Times New Roman
  H1：16pt Bold 宋体 / Times New Roman
  H2：15pt Bold 宋体 / Times New Roman
  H3：14pt Bold 宋体 / Times New Roman
  H4：13pt Bold 宋体 / Times New Roman
  H5：12pt Regular 宋体 / Times New Roman
  正文：12pt Regular 宋体 / Times New Roman

改编自 tender-workflow 的 docx_outline_template.py + docx_writer.py
"""

import os
import re
import sys
import datetime
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    except subprocess.CalledProcessError:
        print("错误：缺少 python-docx 依赖且自动安装失败，请手动执行：", file=sys.stderr)
        print(f"  {sys.executable} -m pip install python-docx", file=sys.stderr)
        sys.exit(1)
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement


# ── 字体设置函数 ──────────────────────────────────

def setup_styles(doc):
    """设置文档样式（字体规范）"""
    def apply_font(style, size_pt, bold=False, cn_font='宋体', latin_font='Times New Roman'):
        style.font.size = Pt(size_pt)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.name = latin_font
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), latin_font)
        rFonts.set(qn('w:hAnsi'), latin_font)
        rFonts.set(qn('w:eastAsia'), cn_font)
        rFonts.set(qn('w:cs'), cn_font)
        for attr in [qn('w:asciiTheme'), qn('w:hAnsiTheme'),
                     qn('w:eastAsiaTheme'), qn('w:cstheme')]:
            rFonts.attrib.pop(attr, None)

    # 字体名称使用 OOXML 标准映射名（跨平台兼容）
    # '宋体': Windows -> SimSun, macOS -> Songti SC
    # '黑体': Windows -> SimHei, macOS -> Heiti SC
    apply_font(doc.styles['Title'],     22, bold=True,  cn_font='黑体')   # 封面：二号黑体加粗
    apply_font(doc.styles['Heading 1'], 16, bold=True,  cn_font='宋体')   # H1：三号宋体加粗
    apply_font(doc.styles['Heading 2'], 15, bold=True,  cn_font='宋体')   # H2：小三宋体加粗
    apply_font(doc.styles['Heading 3'], 14, bold=True,  cn_font='宋体')   # H3：四号宋体加粗

    # H4 和 H5 需要创建或获取样式
    for level, size, bold in [(4, 13, True), (5, 12, False)]:
        style_name = f'Heading {level}'
        try:
            style = doc.styles[style_name]
        except KeyError:
            style = doc.styles.add_style(style_name, 1)  # WD_STYLE_TYPE.PARAGRAPH
            style.base_style = doc.styles['Heading 3']
        apply_font(style, size, bold=bold, cn_font='宋体')

    apply_font(doc.styles['Normal'], 12, bold=False, cn_font='宋体')   # 正文：小四宋体

    # 正文行距 1.5 倍
    doc.styles['Normal'].paragraph_format.line_spacing = 1.5


def setup_heading_numbering(doc):
    """为 Heading 1-5 设置多级列表自动编号（1, 1.1, 1.1.1, ...）"""
    # 确保文档有 numbering part
    try:
        numbering_part = doc.part.numbering_part
    except Exception:
        # 创建 numbering part（python-docx 默认文档可能没有）
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.parts.numbering import NumberingPart
        numbering_part = NumberingPart.new()
        doc.part.relate_to(numbering_part, RT.NUMBERING)

    numbering_elem = numbering_part.element

    # 找到一个不冲突的 abstractNumId
    existing_ids = [
        int(an.get(qn('w:abstractNumId')))
        for an in numbering_elem.findall(qn('w:abstractNum'))
    ]
    abstract_num_id = str(max(existing_ids) + 1) if existing_ids else '0'

    # 创建 abstractNum 定义
    abstract_num = OxmlElement('w:abstractNum')
    abstract_num.set(qn('w:abstractNumId'), abstract_num_id)

    # 多级命名空间标识
    multi_level_type = OxmlElement('w:multiLevelType')
    multi_level_type.set(qn('w:val'), 'multilevel')
    abstract_num.append(multi_level_type)

    # 定义 5 级编号格式
    level_formats = [
        ('%1', '0'),            # H1: "1"
        ('%1.%2', '1'),         # H2: "1.1"
        ('%1.%2.%3', '2'),      # H3: "1.1.1"
        ('%1.%2.%3.%4', '3'),   # H4: "1.1.1.1"
        ('%1.%2.%3.%4.%5', '4'),# H5: "1.1.1.1.1"
    ]

    for lvl_text, ilvl in level_formats:
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), ilvl)

        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)

        num_fmt = OxmlElement('w:numFmt')
        num_fmt.set(qn('w:val'), 'decimal')
        lvl.append(num_fmt)

        lvl_text_elem = OxmlElement('w:lvlText')
        lvl_text_elem.set(qn('w:val'), lvl_text)
        lvl.append(lvl_text_elem)

        lvl_jc = OxmlElement('w:lvlJc')
        lvl_jc.set(qn('w:val'), 'left')
        lvl.append(lvl_jc)

        # 编号后跟空格
        suff = OxmlElement('w:suff')
        suff.set(qn('w:val'), 'space')
        lvl.append(suff)

        abstract_num.append(lvl)

    numbering_elem.append(abstract_num)

    # 创建 num 引用 abstractNum
    existing_num_ids = [
        int(n.get(qn('w:numId')))
        for n in numbering_elem.findall(qn('w:num'))
    ]
    num_id = str(max(existing_num_ids) + 1) if existing_num_ids else '1'

    num_elem = OxmlElement('w:num')
    num_elem.set(qn('w:numId'), num_id)
    abstract_num_id_ref = OxmlElement('w:abstractNumId')
    abstract_num_id_ref.set(qn('w:val'), abstract_num_id)
    num_elem.append(abstract_num_id_ref)
    numbering_elem.append(num_elem)

    # 将 Heading 1-5 样式关联到编号
    for h_level in range(1, 6):
        style_name = f'Heading {h_level}'
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue

        pPr = style.element.get_or_add_pPr()

        # 移除已有的 numPr
        old_numPr = pPr.find(qn('w:numPr'))
        if old_numPr is not None:
            pPr.remove(old_numPr)

        numPr = OxmlElement('w:numPr')
        ilvl_elem = OxmlElement('w:ilvl')
        ilvl_elem.set(qn('w:val'), str(h_level - 1))
        numId_elem = OxmlElement('w:numId')
        numId_elem.set(qn('w:val'), num_id)
        numPr.append(ilvl_elem)
        numPr.append(numId_elem)
        pPr.append(numPr)


def clean_doc_defaults(doc):
    """清除文档默认值中的主题字体引用"""
    styles_elem = doc.styles.element
    doc_defaults = styles_elem.find(qn('w:docDefaults'))
    if doc_defaults is None:
        return
    rPr_default = doc_defaults.find(qn('w:rPrDefault'))
    if rPr_default is None:
        return
    rPr = rPr_default.find(qn('w:rPr'))
    if rPr is None:
        return
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        return
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:cs'), '宋体')
    for attr in [qn('w:asciiTheme'), qn('w:hAnsiTheme'),
                 qn('w:eastAsiaTheme'), qn('w:cstheme')]:
        rFonts.attrib.pop(attr, None)


def apply_run_font(run, cn_font='宋体', latin_font='Times New Roman'):
    """对单个 run 设置字体（清除主题字体继承）"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), latin_font)
    rFonts.set(qn('w:hAnsi'), latin_font)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:cs'), cn_font)
    for attr in [qn('w:asciiTheme'), qn('w:hAnsiTheme'),
                 qn('w:eastAsiaTheme'), qn('w:cstheme')]:
        rFonts.attrib.pop(attr, None)


# ── 包装函数 ──────────────────────────────────────

def add_heading_cn(doc, text, level, cn_font='宋体'):
    """添加中文标题"""
    font = '黑体' if level == 0 else cn_font
    h = doc.add_heading(text, level)
    for run in h.runs:
        apply_run_font(run, cn_font=font)
    return h


def add_para_cn(doc, text, cn_font='宋体', bold=False, italic=False, align=None):
    """添加中文段落"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        apply_run_font(run, cn_font=cn_font)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
    if align:
        p.alignment = align
    return p


def add_picture_cn(doc, img_path, caption=None, width_cm=14.0, cn_font='宋体'):
    """添加图片（带描述文字），文件不存在或格式异常时降级为占位符"""
    try:
        doc.add_picture(img_path, width=Cm(width_cm))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        add_para_cn(doc, f'[图片占位符 — {caption or "无描述"}（路径: {img_path}）：请手动插入图片]',
                    cn_font=cn_font, italic=True)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if caption:
        add_para_cn(doc, caption, cn_font=cn_font, italic=True)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


# ── Markdown 渲染 ──────────────────────────────────

def _apply_inline_formatting(paragraph, text, cn_font='宋体'):
    """处理行内格式：**bold**、*italic*、`code`"""
    # 简单的正则分割处理
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
        else:
            run = paragraph.add_run(part)
        apply_run_font(run, cn_font=cn_font)


def write_markdown(doc, md_text, cn_font='宋体', base_dir=None):
    """将 Markdown 文本渲染到 DOCX 文档

    Args:
        base_dir: 图片相对路径的基准目录。如果为 None，使用当前工作目录。
    """
    lines = md_text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 空行跳过
        if not stripped:
            i += 1
            continue

        # 标题（要求 # 后有空格，避免误匹配 #hashtag）
        heading_match = re.match(r'^(#{1,5})\s+(.+)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            title_text = heading_match.group(2).strip()
            add_heading_cn(doc, title_text, level, cn_font=cn_font)
            i += 1
            continue

        # 图片
        img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)', stripped)
        if img_match:
            caption = img_match.group(1)
            img_path = img_match.group(2)
            # 将相对路径解析为绝对路径，按优先级尝试多个 base_dir
            if not os.path.isabs(img_path):
                resolved = None
                candidates = []
                if base_dir:
                    candidates.append(os.path.join(base_dir, img_path))
                candidates.append(os.path.join(os.getcwd(), img_path))
                for candidate in candidates:
                    if os.path.isfile(candidate):
                        resolved = candidate
                        break
                img_path = resolved or candidates[0]
            add_picture_cn(doc, img_path, caption=caption, cn_font=cn_font)
            i += 1
            continue

        # 表格
        if '|' in stripped and stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            _render_table(doc, table_lines, cn_font)
            continue

        # 引用
        if stripped.startswith('>'):
            quote_text = stripped.lstrip('>').strip()
            p = add_para_cn(doc, quote_text, cn_font=cn_font, italic=True)
            p.paragraph_format.left_indent = Cm(0.75)
            i += 1
            continue

        # 无序列表
        if stripped.startswith(('- ', '* ')):
            p = doc.add_paragraph(style='List Bullet')
            _apply_inline_formatting(p, stripped[2:], cn_font)
            i += 1
            continue

        # 有序列表
        ol_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if ol_match:
            p = doc.add_paragraph(style='List Number')
            _apply_inline_formatting(p, ol_match.group(2), cn_font)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        _apply_inline_formatting(p, stripped, cn_font)
        i += 1


def _render_table(doc, table_lines, cn_font='宋体'):
    """渲染 Markdown 表格"""
    # 过滤分隔行
    data_lines = [l for l in table_lines if not re.match(r'^\|[\s\-:]+\|$', l)]
    if not data_lines:
        return

    rows = []
    for line in data_lines:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < num_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = cell_text
                for p in cell.paragraphs:
                    for run in p.runs:
                        apply_run_font(run, cn_font=cn_font)


# ── 文档创建 ──────────────────────────────────────

def create_document(
    title: str = '解决方案',
    subtitle: str = '',
    author: str = '',
) -> Document:
    """创建一个配置好样式的空白文档"""
    doc = Document()

    # 页面设置（A4）
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.4)
    sec.right_margin = Cm(2.4)

    setup_styles(doc)
    clean_doc_defaults(doc)
    setup_heading_numbering(doc)

    # 封面
    add_heading_cn(doc, title, 0, cn_font='黑体')
    if subtitle:
        add_para_cn(doc, subtitle)
    add_para_cn(doc, f'生成时间：{datetime.date.today()}')
    if author:
        add_para_cn(doc, author)
    doc.add_page_break()

    # 目录占位
    add_heading_cn(doc, '目录', 1)
    add_para_cn(doc, '[Word 自动生成目录，按 Ctrl+A -> F9 更新]')
    doc.add_page_break()

    return doc


def save_document(doc: Document, output_path: str) -> str:
    """保存文档"""
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


# ── CLI 入口 ──────────────────────────────────────

def main():
    import argparse
    parser = argparse.ArgumentParser(description='Solution Master DOCX 渲染器')
    parser.add_argument('input', help='输入 Markdown 文件路径')
    parser.add_argument('--output', '-o', help='输出 DOCX 文件路径')
    parser.add_argument('--title', default='解决方案', help='文档标题')
    parser.add_argument('--subtitle', default='', help='副标题')
    parser.add_argument('--author', default='', help='作者')
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f'错误：文件不存在: {input_path}', file=sys.stderr)
        sys.exit(1)

    md_text = input_path.read_text(encoding='utf-8')

    output_path = args.output or str(input_path.with_suffix('.docx'))

    doc = create_document(title=args.title, subtitle=args.subtitle, author=args.author)
    write_markdown(doc, md_text, base_dir=str(input_path.resolve().parent))
    save_document(doc, output_path)

    print(f'DOCX 生成完成: {output_path}')


if __name__ == '__main__':
    main()
